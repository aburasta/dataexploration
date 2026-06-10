#!/usr/bin/env python3
"""Overall Production Guide (Word .docx) for the finance x animal-behavior
animated documentary series. Folds in the visual style module + the full
per-episode workflow. Designed to be uploaded to the Claude Projects section."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FULL_SUFFIX = (
    "Bright, cheerful, wholesome flat 2D vector children's-book clip-art illustration. Rounded, cute, friendly "
    "characters and objects with clean simple shapes and chunky proportions. Soft subtle shading and gentle soft "
    "highlights for a little depth; thin minimal darker edge accents only where a shape needs separation (no heavy "
    "black outlines). Small dot eyes, rosy cheek blush, simple happy expressions. Bright multicolor brand palette: "
    "sunshine yellow (#FBC02D), sky blue (#4FC3E8), leaf green (#7CC576), tomato coral (#F0654E), tangerine "
    "(#F39237), blush pink (#F6A9B6), stone gray (#9DB0BC), warm brown (#8D6E63), cream (#FBF7EF), charcoal "
    "(#3A3631). SINGLE isolated element only, centered, on a plain TRANSPARENT background — no scene, no setting, no "
    "ground shadow. Clean, modern, wholesome. No grain, no gradients, no 3D, no realistic textures, no photorealism, "
    "no heavy black outlines, no clutter, no scene background."
)
COMPACT_SUFFIX = (
    "STYLE: bright cheerful flat 2D vector children's-book clip-art; rounded cute friendly shapes, soft subtle "
    "shading + gentle highlights, thin minimal edge accents (no heavy black outlines); dot eyes, rosy cheeks; "
    "palette sunshine #FBC02D, sky #4FC3E8, leaf #7CC576, coral #F0654E, tangerine #F39237, blush #F6A9B6, stone "
    "#9DB0BC, brown #8D6E63, cream #FBF7EF, charcoal #3A3631; SINGLE isolated element, centered, TRANSPARENT "
    "background, no scene/no ground shadow; clean modern wholesome. No grain/gradient/3D/realistic/photoreal/heavy "
    "outline/clutter/scene-background."
)
NEGATIVE = (
    "scene background, environment, ground shadow, multiple subjects, photorealism, 3D render, CGI, realistic "
    "textures, film grain, gradients banding, glossy/metallic highlights, ambient occlusion, drop shadows, lens "
    "flare, glow/bloom, heavy black outlines, comic linework, sketchy wobble, crosshatching, dark/gritty mood, "
    "cluttered detail, complex detailed faces, tiny illegible text, watermarks, extra limbs/fingers, natural "
    "photographic animal coloring."
)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def h1(doc, text):
    return doc.add_heading(text, level=1)


def h2(doc, text):
    return doc.add_heading(text, level=2)


def body(doc, text, size=10.5, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(size)
    return p


def numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(size)
    return p


def box(doc, label, text):
    if label:
        lp = doc.add_paragraph()
        r = lp.add_run(label)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("1F4E5F")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F2F4F0")
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


def callout(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, "FBEFD6")
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("8A4B16")
    doc.add_paragraph()


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ============================ TITLE ============================
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PRODUCTION GUIDE"); r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string("1F4E5F")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Finance × Animal Behavior — Animated YouTube Documentary Series")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string("EF5A2A")
s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s2.add_run("Master instructions for turning ONE row of the concept spreadsheet into a finished episode "
               "(shot list + transparent element PNGs assembled in VideoScribe). Upload to the Claude Projects section.")
r.italic = True; r.font.size = Pt(10)
doc.add_paragraph()

# ============================ 0. MISSION ============================
h1(doc, "0. What This Project Does")
body(doc, "This is a YouTube documentary series that explains difficult finance concepts by paralleling them to "
          "specific animal behaviors. Each episode takes ONE row from the file "
          "“finance_animal_behavior_matches.xlsx” and turns it into a short, highly engaging animated explainer.")
body(doc, "When the user names or pastes a row (a finance concept + its animal-behavior match), Claude must produce "
          "the deliverables defined in Part 2 — a shot list plus a folder of transparent element PNGs assembled in "
          "VideoScribe — following every rule in this guide. The visual style (Part 7) and "
          "the script/hook formula (Part 4) are FIXED and identical for every episode — only the content changes.")

callout(doc, "THE ONE RULE ABOVE ALL: Every single video uses the exact same visual style defined in Part 7. "
             "The style never changes from episode to episode. Consistency across episodes IS the brand.")

# ============================ 1. INPUT ============================
h1(doc, "1. The Input: One Spreadsheet Row")
body(doc, "Source file: finance_animal_behavior_matches.xlsx. Each row contains:")
bullet(doc, "Category, Granular Concept, and The Specific Finance Aspect — the precise idea to teach.")
bullet(doc, "Animal Match 1 & 2 (name + a script-ready behavior description + a quality %).")
bullet(doc, "Notes / verification flags.")
body(doc, "Default behavior: build the episode around the HIGHER-rated match. Use the second match only if it adds a "
          "distinct facet, or in the honesty beat (see Part 4). If a match is flagged below 75% or “debated,” lean "
          "into the “where the analogy breaks down” section rather than overstating it.")

# ============================ 2. OUTPUT OVERVIEW ============================
h1(doc, "2. The Output: Three Deliverables, Every Time")
body(doc, "For each row called, Claude produces three things:")
numbered(doc, "DELIVERABLE A — THE SHOT LIST (a Word document). A reader-ready narration script broken into FRAMES, "
              "each with a timestamp range, the on-screen visual, and the breakdown of which isolated elements are "
              "drawn on during that stretch of script. Visuals change every 3–5 seconds. (Spec: Part 3.)")
numbered(doc, "DELIVERABLE B — THE HIGGSFIELD PROMPT PACK. A single, paste-ready set of per-ELEMENT prompts grouped "
              "by frame that generate EVERY element for the episode in the fixed style, written with thorough, "
              "repeated descriptions so the whole set stays consistent. The user pastes these into the Higgsfield web "
              "app (Seedream 4.5) to generate the elements for free. (Spec: Part 8f.)")
numbered(doc, "DELIVERABLE C — THE FRAMES FOLDER. For each frame, every element as a TRANSPARENT PNG, produced by "
              "running the Prompt Pack + background removal — generated free in the Higgsfield web app, or via API at "
              "credit cost — named to match the frames. (Spec: Part 9.)")
body(doc, "All three are linked: Frame 001 in the document ↔ its element prompts in the Pack ↔ the element PNGs in "
          "frames/frame_001/ in the folder.")

# ============================ 3. DELIVERABLE A SPEC ============================
h1(doc, "3. Deliverable A — The Shot List Document")
body(doc, "Structure of the document, top to bottom:")
numbered(doc, "Header block: Episode title, the finance concept, the animal match used, target runtime, total word "
              "count, and narration pace used (default 150 words/min).")
numbered(doc, "CLEAN NARRATION SCRIPT: the full voice-over written as continuous prose, ready to read aloud "
              "start to finish with no interruptions. This is what the narrator records.")
numbered(doc, "THE SHOT LIST TABLE: the same script broken into timed FRAMES with element instructions (columns below).")
body(doc, "Shot list table columns:")
bullet(doc, "Frame # (e.g., 001).")
bullet(doc, "Timestamp (start–end, mm:ss–mm:ss).")
bullet(doc, "Narration for this frame (the exact words spoken during this stretch — ~8–13 words = 3–5 seconds).")
bullet(doc, "On-screen visual (one-line description of what the finished frame shows).")
bullet(doc, "Element breakdown + draw order (each isolated element in the frame, listed in the order the hand draws "
            "them on — e.g. 1) host fox, 2) coral droplet).")
bullet(doc, "Prompt #s (the per-element prompt numbers in the Prompt Pack that generate this frame's elements). The "
            "FULL prompts live in the Pack — not the shot list — so they stay together and consistent.")
bullet(doc, "Type (STILL or CLIP — clip only where real motion is essential; default motion is the VideoScribe "
            "hand-reveal, see Part 8).")

# ============================ 4. SCRIPT RULES ============================
h1(doc, "4. Script Rules (identical every episode)")

h2(doc, "4a. The Hook (first ~5–15 seconds) — the most important part")
body(doc, "The hook ALWAYS creates suspense around the personal BENEFIT of knowing the lesson. Frame it as an "
          "intriguing question (or a “what if”) that makes the viewer feel they’re about to gain an edge most people "
          "miss. Do NOT open by naming the concept; open with the payoff and the mystery.")
body(doc, "Formula: [Tease the benefit / the edge] → [pose a question or “what if” the viewer wants answered] → "
          "[promise the animal will reveal the answer]. Keep it punchy.")
box(doc, "HOOK EXAMPLES (benefit-led, suspenseful):",
    "Lesson = avoid herd behavior →\n"
    "  “How do the people who spot the next big thing always see it before everyone else? "
    "The answer is hiding in how a single bird breaks from the flock...”\n\n"
    "Lesson = risk pooling / insurance →\n"
    "  “What if one bad month could never wipe you out — no matter how unlucky you got? "
    "A tiny vampire bat has quietly solved this for millions of years.”\n\n"
    "Lesson = compounding →\n"
    "  “What if doing one small, boring thing today decided whether you’re rich in 20 years? "
    "A colony of bees shows exactly when the magic kicks in.”")

h2(doc, "4b. Retention Structure (after the hook)")
body(doc, "Arrange the body to hold attention as tightly as possible:")
numbered(doc, "HOOK — benefit + question (above).")
numbered(doc, "STAKES — quickly make the viewer feel why getting this wrong hurts them (loss aversion pulls attention).")
numbered(doc, "MEET THE ANIMAL — introduce the animal and its behavior vividly (use the spreadsheet description).")
numbered(doc, "THE PARALLEL — map the behavior onto the finance aspect, step by step, one idea at a time.")
numbered(doc, "THE TURN / “WHERE IT BREAKS DOWN” — honestly note where the analogy stops (this builds trust and "
              "teaches nuance; it is a signature of the series).")
numbered(doc, "THE PAYOFF — deliver the practical lesson promised in the hook: what the viewer should DO or notice.")
numbered(doc, "BUTTON — a short, memorable closing line (and a soft prompt to subscribe / next episode).")
body(doc, "Throughout: short sentences, concrete imagery, open loops (tease something resolved later), and a question "
          "every ~20–30 seconds to re-grab attention. Conversational, warm, smart — never lecture-y.")

h2(doc, "4c. Length & Pace")
bullet(doc, "Target runtime: 10–15 minutes (default ~12). State the chosen target in the header.")
bullet(doc, "Narration pace: 150 words/minute = 2.5 words/second (documentary-calm). Adjustable, but state it.")
bullet(doc, "At 150 wpm: a 10-minute episode ≈ 1,500 words, 12 min ≈ 1,800 words, 15 min ≈ 2,250 words.")

# ============================ 5. TIMING METHOD ============================
h1(doc, "5. Timing & Timestamp Method (every frame = 3–5 seconds)")
body(doc, "How Claude assigns timestamps and the per-frame visual cadence:")
numbered(doc, "Write the full clean narration script first.")
numbered(doc, "Compute total duration: total_words ÷ 2.5 = seconds (at 150 wpm).")
numbered(doc, "Segment the script into FRAMES of 3–5 seconds each. At 2.5 words/sec that is roughly 8–13 words per "
              "frame — usually one short sentence or clause per frame.")
numbered(doc, "Assign each frame a cumulative timestamp range (frame 1: 00:00–00:04, frame 2: 00:04–00:09, etc.), "
              "computing each frame’s length as its_word_count ÷ 2.5.")
numbered(doc, "List each frame's ELEMENTS and the order the hand draws them; the frame's elements are drawn on in "
              "sequence (draw order) across its narration slice. This guarantees the on-screen visual changes "
              "every 3–5 seconds, as required.")
callout(doc, "RESULT: a 10–15 minute episode produces roughly 150–300 FRAMES, each composed of a few isolated "
             "element PNGs drawn on in sequence. Reuse saved ELEMENTS (Host, recurring props) across frames rather "
             "than regenerating the same asset every time (see Part 8e).")
body(doc, "Round timestamps to whole seconds. Keep the sum of frame lengths equal to the total duration.")

# ============================ 6. SHOT LIST EXAMPLE ============================
h1(doc, "6. Worked Example — Shot List Excerpt (Frames)")
body(doc, "Concept: Risk pooling. Animal: vampire bat reciprocal blood sharing (96%). The first ~25 seconds, to show "
          "the exact format Claude must produce — each frame broken into the isolated elements the hand draws on, in "
          "order (per-element prompts live in the Prompt Pack, Part 8f).")

ex = [
    ("001", "00:00–00:05", "What if one bad month could never wipe you out — no matter how unlucky you got?",
     "Host sky-blue fox leaning in, curious, rosy cheeks",
     "1) host fox (sky blue, coral scarf); 2) small coral question-mark", "el01–el02"),
    ("002", "00:05–00:09", "Most people are just one emergency away from losing everything.",
     "A lone stone-gray figure standing small; a thin coral crack",
     "1) lone gray figure; 2) coral crack motif", "el01–el02"),
    ("003", "00:09–00:13", "But some never seem to get destroyed by a single disaster.",
     "Three rounded figures together, calm, a soft leaf-green protective dome",
     "1) three figures; 2) leaf-green dome", "el01–el02"),
    ("004", "00:13–00:18", "Their secret was solved millions of years ago — by a tiny vampire bat.",
     "Cute stone-gray + sky-blue vampire bat hanging, friendly dot eyes, rosy cheeks",
     "1) vampire bat element", "el01"),
    ("005", "00:18–00:23", "A vampire bat must drink within two days, or it dies.",
     "Same bat; a small sunshine moon + clock motif",
     "1) vampire bat (reuse el); 2) moon/clock motif", "el01–el02"),
    ("006", "00:23–00:27", "And on any night, some of them fail to feed.",
     "Two bats: one with a coral droplet (fed), one empty and worried",
     "1) fed bat + coral droplet; 2) hungry bat", "el01–el02"),
]
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Frame", "Timestamp", "Narration", "On-screen visual", "Element breakdown + draw order",
                       "Prompt #s"]):
    set_cell_text(hdr[i], h, bold=True, size=8.5, color="FFFFFF")
    shade_cell(hdr[i], "1F4E5F")
for frame, ts, narr, vis, els, prm in ex:
    c = tbl.add_row().cells
    set_cell_text(c[0], frame, size=8.5)
    set_cell_text(c[1], ts, size=8.5)
    set_cell_text(c[2], narr, size=8.5)
    set_cell_text(c[3], vis, size=8.5)
    set_cell_text(c[4], els, size=8.5)
    set_cell_text(c[5], prm, size=8.5)
for row in tbl.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(0.8)
    row.cells[2].width = Inches(1.9)
    row.cells[3].width = Inches(1.7)
    row.cells[4].width = Inches(1.7)
    row.cells[5].width = Inches(0.5)
doc.add_paragraph()
body(doc, "Note how each frame is one short clause (~3–5s) decomposed into a few isolated transparent element PNGs, "
          "drawn on in order by the hand, and how elements recur (the same bat is reused as a saved Element) so the "
          "frames flow as continuous animation. The full document would also include the clean narration script "
          "above this table.", italic=True)

# ============================ 7. VISUAL STYLE MODULE ============================
h1(doc, "7. Visual Style — Applied to EVERY Episode")
callout(doc, "This style is mandatory and unchanging. Every element in every episode must match it. Append the Master "
             "Style Suffix (Part 8) to every Higgsfield prompt. Every element is generated ISOLATED on a transparent "
             "background.")

body(doc, "The look in one sentence: bright, cheerful, wholesome FLAT 2D vector children's-book CLIP-ART — rounded "
          "cute friendly characters and objects, clean simple shapes, soft subtle shading with gentle highlights, "
          "thin minimal edge accents (no heavy black outlines), dot eyes and rosy cheeks; every asset a SINGLE "
          "isolated element on a transparent background.")

h2(doc, "7a. Core Traits")
for tr in [
    "CLIP-ART, ISOLATED — one element per asset, centered, on a TRANSPARENT background (no scene, no ground shadow).",
    "ROUNDED & CUTE — chunky friendly proportions; circles, ovals, soft rounded shapes.",
    "SOFT SHADING + HIGHLIGHTS — gentle subtle shading and soft highlights for a little depth (not flat-only, not 3D).",
    "THIN MINIMAL EDGE ACCENTS — light darker edges only where shapes need separation; NO heavy black ink outlines.",
    "BRIGHT MULTICOLOR BRAND PALETTE (below) — cheerful and saturated, but fixed for series consistency.",
    "FRIENDLY FACES — small dot eyes, rosy cheek blush, simple happy expressions.",
    "CLEAN & WHOLESOME — uncluttered, modern, no grain, no texture overlays.",
]:
    bullet(doc, tr)

h2(doc, "7b. Color Palette (canonical brand hex)")
palette = [
    ("Sunshine Yellow", "#FBC02D", "Primary", "Beaks, manes, warm highlights"),
    ("Sky Blue", "#4FC3E8", "Secondary", "Water/sky elements, calm characters"),
    ("Leaf Green", "#7CC576", "Secondary", "Foliage, ground shapes"),
    ("Tomato Coral", "#F0654E", "Accent", "THE focal color — the key idea (sparingly)"),
    ("Tangerine", "#F39237", "Accent", "Secondary warm accent (lion/monkey)"),
    ("Blush Pink", "#F6A9B6", "Accent", "Cheek blush, soft accents"),
    ("Stone Gray", "#9DB0BC", "Neutral", "Elephant/rhino bodies, neutral props"),
    ("Warm Brown", "#8D6E63", "Neutral", "Trunks, wood, branches"),
    ("Cream", "#FBF7EF", "Primary", "VideoScribe canvas / negative space"),
    ("Charcoal", "#3A3631", "Neutral", "Eyes / edge accents / fine detail — never pure black"),
]
pt = doc.add_table(rows=1, cols=5); pt.style = "Table Grid"
hd = pt.rows[0].cells
for i, x in enumerate(["Swatch", "Name", "Hex", "Role", "Usage"]):
    set_cell_text(hd[i], x, bold=True, size=8.5, color="FFFFFF"); shade_cell(hd[i], "1F4E5F")
for name, hexv, role, usage in palette:
    c = pt.add_row().cells
    shade_cell(c[0], hexv); set_cell_text(c[0], "", size=8.5)
    set_cell_text(c[1], name, bold=True, size=8.5)
    set_cell_text(c[2], hexv, size=8.5)
    set_cell_text(c[3], role, size=8.5)
    set_cell_text(c[4], usage, size=8.5)
for row in pt.rows:
    row.cells[0].width = Inches(0.5)
    row.cells[1].width = Inches(1.3)
    row.cells[2].width = Inches(0.9)
    row.cells[3].width = Inches(0.9)
    row.cells[4].width = Inches(2.8)
doc.add_paragraph()
bullet(doc, "Bright but balanced: use the palette across elements; reserve TOMATO CORAL for the one focal element per "
            "frame (the key idea). Cream is the VideoScribe canvas. No pure black, no pure white.")

h2(doc, "7c. Characters")
bullet(doc, "THE HOST (recurring narrator): one fixed friendly clip-art animal guide — recommended a rounded sky-blue "
            "fox with a coral scarf, dot eyes and rosy cheeks — appearing in every episode. Lock ONE design and reuse "
            "it identically as a saved Element.")
bullet(doc, "ANIMAL OF THE WEEK: the spreadsheet animal as a cute rounded clip-art character, RE-TINTED into the "
            "brand palette (NOT natural photo colors), simplified into friendly rounded shapes.")
bullet(doc, "CROWDS / MARKETS / “YOU”: simple rounded clip-art figures in palette colors; the odd-one-out is coral.")

h2(doc, "7d. Finance Motifs (in-style)")
bullet(doc, "Charts: chunky rounded bars/lines, key line in coral, minimal axes — never realistic stock screenshots. "
            "Each is its OWN isolated transparent element.")
bullet(doc, "Money: cute chunky coins; stacks as rounded rectangles. Arrows: thick rounded; coral/tangerine for "
            "alarm. Each drawn as a single isolated transparent element.")
bullet(doc, "Keep one consistent icon vocabulary across all episodes (herd = identical shapes, growth = sprouting "
            "plant, etc.).")

h2(doc, "7e. Composition")
bullet(doc, "Elements are individual assets — NO fixed aspect ratio on the PNG. The 16:9 LAYOUT is composed on the "
            "VideoScribe CREAM canvas, not on the PNGs. Highest available resolution per element.")
bullet(doc, "One clear focal subject per frame; arrange elements with breathing room on the canvas.")
bullet(doc, "Keep the bottom ~15% of the canvas clear (caption-safe). Vary element scale across a sequence.")

# ============================ 8. IMAGE GENERATION ============================
h1(doc, "8. Image Generation with Higgsfield")

h2(doc, "8a. Model & Settings")
bullet(doc, "PRIMARY model — Seedream 4.5 (model id: seedream_v4_5). This is the DEFAULT for every element in the "
            "series: high quality, up to 4K, ~1 credit each (half the cost of Nano Banana Pro), and included in the "
            "paid plans’ 365-day UNLIMITED set. Use quality 'high' for hero elements, 'basic' for the rest.")
bullet(doc, "SECONDARY model — Nano Banana Pro (nano_banana_pro) or GPT Image, used ONLY for elements that contain "
            "text or numbers (stronger text rendering). These cost more credits, so use them sparingly.")
bullet(doc, "Generate each element as a SINGLE isolated subject on a plain flat background — NO per-PNG 16:9 framing; "
            "the 16:9 layout is composed later on the VideoScribe canvas. Generate one element per prompt.")
bullet(doc, "Then REMOVE THE BACKGROUND to get a transparent-background PNG (the alpha cutout) — see Part 8g.")
bullet(doc, "For optional motion shots, generate the element still first (Seedream), then animate it with "
            "image-to-video so the style is preserved.")
bullet(doc, "COST NOTE: the Seedream ‘365-day Unlimited’ benefit is advertised for the Higgsfield WEB APP — "
            "generating there avoids burning credits. Via API/automation Seedream still bills ~1 credit/image, so for "
            "daily volume either do the bulk in the web app or budget credits. The FREE tier (~10 credits/day) "
            "cannot sustain a full episode.")

h2(doc, "8b. Consistency (so every element stays on-model)")
bullet(doc, "Save the Host (and recurring props) as a Higgsfield ELEMENT and reference it in every prompt so the "
            "character is identical across all frames and all episodes. Reuse saved Elements rather than regenerating.")
bullet(doc, "Keep one approved “style key-frame” and pass it as a reference image so new elements inherit "
            "palette/finish.")
bullet(doc, "Per element, reuse ONE SEED per character/prop so the same asset looks identical in every frame it "
            "appears in.")
bullet(doc, "Describe recurring characters with the SAME wording every time — never paraphrase or write 'same as "
            "before' (paraphrasing causes drift).")

h2(doc, "8c. When to use a CLIP instead of the hand-reveal")
body(doc, "Default motion is the VideoScribe hand-reveal (the hand draws each element on). Use a short (2–4s) "
          "image-to-video CLIP only as an OPTIONAL motion path where real motion is essential and cannot be faked by "
          "the hand-reveal, e.g.: a number ticking up, a chart drawing itself, a coin flip, or a transformation. "
          "Generate the element still first, then animate it. Mark these as CLIP in the shot list.")

h2(doc, "8d. The Master Prompt Formula")
body(doc, "Every element prompt = [the ONE isolated element described identically] + [single clear pose/state] + "
          "[where the coral accent goes] + 'single isolated element, centered, transparent background, no scene, no "
          "ground shadow' + the Master Style Suffix appended verbatim. Each prompt is ONE isolated element on a "
          "transparent background.")
box(doc, "MASTER STYLE SUFFIX (append to EVERY element prompt):", FULL_SUFFIX)
box(doc, "NEGATIVE / AVOID:", NEGATIVE)

h2(doc, "8e. Managing Credit Cost (important at 10–15 minutes)")
body(doc, "A 10–15 minute episode with a visual change every 3–5 seconds implies ~150–300 frames, each made of a few "
          "elements. Generating a brand-new element for every appearance is the most credit-expensive path. To keep "
          "the pacing but control cost:")
bullet(doc, "REUSE SAVED ELEMENTS across frames — generate the Host, recurring props, and the animal of the week ONCE "
            "as Elements, then place the same transparent PNG on the canvas wherever it recurs instead of "
            "regenerating it. Most frames reuse existing elements and only add one or two new ones.")
bullet(doc, "Generate the BULK FOR FREE in the Higgsfield WEB APP on Seedream 4.5 (unlimited on a paid plan), rather "
            "than via API where each image bills ~1 credit.")
bullet(doc, "Reserve true CLIPs and re-generations for where they truly matter — they cost the most.")
bullet(doc, "Cost guide: Seedream 4.5 ≈ 1 credit/image via API (or UNLIMITED in the web app on a paid plan); Nano "
            "Banana Pro ≈ 2 credits/image. Reusing saved Elements means you generate only the NEW elements per frame, "
            "and the web app’s unlimited Seedream removes the credit ceiling for elements entirely.")

h2(doc, "8f. Deliverable B — The Higgsfield Prompt Pack")
body(doc, "After the shot list is built, ALWAYS output a PROMPT PACK: the exact, paste-ready per-ELEMENT prompts the "
          "user runs in the Higgsfield web app (Seedream 4.5) to generate the whole episode's elements consistently. "
          "Each frame becomes a small group of element prompts. It has two parts.")
body(doc, "PART 1 — THE CONSISTENCY HEADER (written once; applies to every element). A fixed block the user keeps "
          "attached to every generation, containing:")
bullet(doc, "STYLE BLOCK: the COMPACT style tag (below), verbatim.")
bullet(doc, "CHARACTER BIBLE: a thorough, word-for-word description of every recurring character that must look "
            "identical in every frame — the Host (species, exact body color, face, signature accessory), the Animal "
            "of the Week (palette re-tint + shape), and the crowd/figure convention. These EXACT words are reused in "
            "every prompt that features the character.")
bullet(doc, "PROP / MOTIF BIBLE: the recurring props and finance motifs (coins, charts, arrows) described identically.")
bullet(doc, "PALETTE: the canonical hex list (Part 7b).")
bullet(doc, "SETTINGS: Seedream 4.5; attach the approved style key-frame as a reference image and the Host Element; "
            "reuse ONE seed per character/prop; generate each as an ISOLATED element on a plain background → remove "
            "background to a transparent PNG.")
body(doc, "PART 2 — THE PER-ELEMENT PROMPTS (grouped by frame, numbered to match the shot list). Group prompts under "
          "each Frame # and list the DRAW ORDER (the order the hand reveals them on the canvas). Each prompt = the "
          "Consistency Header content + [the ONE isolated element] + [coral accent if focal]. Each per-element prompt "
          "MUST:")
bullet(doc, "Be ONE isolated element — a character, a prop, or a motif — asking for a 'single isolated element, "
            "centered, transparent background, no scene, no ground shadow'.")
bullet(doc, "Re-state the recurring characters using the SAME wording as the Character Bible — never paraphrase and "
            "never write 'same as before'; paraphrasing is what causes drift.")
bullet(doc, "Reuse a saved ELEMENT instead of re-prompting the Host/recurring props where possible; state the seed.")
bullet(doc, "Be fully self-contained and paste-ready, so it still works if pasted on its own.")
bullet(doc, "Stay UNDER 3000 characters (Higgsfield's per-prompt limit). Use the COMPACT STYLE TAG below — NOT the "
            "full style paragraph — so each self-contained prompt lands ~500–900 characters.")
box(doc, "COMPACT STYLE TAG (append to every per-element prompt, to stay under 3000 chars):", COMPACT_SUFFIX)
callout(doc, "THOROUGHNESS RULE: every per-element prompt must carry a COMPLETE description of the element, palette, "
             "and style, and ask for a single isolated element on a transparent background — that is what keeps every "
             "separately-generated asset on-model and clean to layer in VideoScribe.")
box(doc, "PROMPT PACK — EXAMPLE (a frame with 2 isolated element prompts):",
    "=== CONSISTENCY HEADER (attach to every element) ===\n"
    "HOST: a friendly rounded sky-blue fox, small charcoal dot eyes, rosy cheeks, a signature coral knitted scarf, "
    "chunky rounded clip-art shapes.\n"
    "STYLE: [paste compact style tag].  SETTINGS: Seedream 4.5, reference = style key-frame, Element = host-fox, "
    "seed per character/prop, isolated element on plain background → remove background to transparent PNG.\n\n"
    "=== FRAME 002 (draw order: el01 then el02) ===\n"
    "el01 — Single isolated clip-art figure: a small lone stone-gray rounded person, dot eyes, rosy cheeks, worried, "
    "centered. Single isolated element, transparent background, no scene, no ground shadow. [compact style tag]\n\n"
    "el02 — Single isolated clip-art motif: a thin tomato-coral jagged crack shape (the focal accent), centered. "
    "Single isolated element, transparent background, no scene, no ground shadow. [compact style tag]")
doc.add_paragraph()

# ---- 8g + 8h: Higgsfield execution + VideoScribe assembly notes ----
h2(doc, "8g. Making the transparent PNG (generate → remove background)")
numbered(doc, "Generate the isolated element on a plain flat background in Higgsfield (Seedream 4.5 for the bulk; Nano "
              "Banana Pro for text elements). Keep 2–4 variations and pick the cleanest, most on-model one.")
numbered(doc, "Run BACKGROUND REMOVAL on the chosen image: in the web app use the remove-background tool, or via MCP "
              "call remove_background with the generation's media_id and media_type='image'. Result = an alpha cutout.")
numbered(doc, "Download the cutout and save it as frames/frame_0NN/frame_0NN_elXX_<name>.png (PNG with alpha). Reuse "
              "saved Host/props as ELEMENTS rather than regenerating.")
bullet(doc, "OPTIONAL SVG line version: auto-trace the PNG (Illustrator Image Trace / Inkscape) so it draws as true "
            "pen strokes in VideoScribe; flat PNGs reveal under the hand.")

h2(doc, "8h. Assembling in VideoScribe (the hand-reveal)")
numbered(doc, "New scribe → set the canvas to 16:9 and the background to CREAM (#FBF7EF) — the brand canvas. Import the "
              "VO audio first so element draw-times match the narration.")
numbered(doc, "For each frame, import that frame's element PNGs from frames/frame_0NN/; position and scale each into "
              "the 16:9 layout (bottom ~15% caption-safe).")
numbered(doc, "Per element, set DRAW TIME / PAUSE / HAND so they reveal in the planned DRAW ORDER across the frame's "
              "narration slice (frame total ≈ 3–5s). Use the same hand throughout.")
numbered(doc, "Order all elements of the frame, move to the next frame, repeat; add captions and music. Drop any CLIP "
              "MP4s onto the timeline at their frame. Render/Download → MP4 at 1080p+. That render is the episode.")

# ============================ 9. FILE OUTPUT ============================
h1(doc, "9. Output Files & Naming")
body(doc, "Produce this structure per episode (EP number + short concept slug):")
box(doc, "FOLDER STRUCTURE:",
    "EP07_risk-pooling_vampire-bat/\n"
    "    EP07_shot-list.docx            <- Deliverable A\n"
    "    EP07_prompt-pack.txt           <- Deliverable B\n"
    "    frames/                        <- Deliverable C (transparent element PNGs)\n"
    "        frame_001/  frame_001_el01_<name>.png  frame_001_el02_<name>.png ...\n"
    "        frame_002/  ...\n"
    "    videoscribe/  EP07.scribe      <- the VideoScribe assembly project\n"
    "    clips/   frame_0XX.mp4         <- OPTIONAL image-to-video for special motion")
bullet(doc, "Numbers MUST match: a frame's element prompts ↔ Frame 0NN ↔ frames/frame_0NN/ element PNGs. Zero-pad "
            "frame numbers to 3 digits and element numbers to 2 digits (frame_007_el02_<name>.png).")
bullet(doc, "Each frame gets a folder of its isolated element PNGs. Optional CLIPs go in /clips with the same frame "
            "number and are marked CLIP; the VideoScribe project lives in /videoscribe.")

# ============================ 10. PROCEDURE ============================
h1(doc, "10. Step-by-Step Procedure When a Row Is Called")
for step in [
    "Read the row: note the concept, the specific finance aspect, the chosen animal match (higher %), and any flags.",
    "Write the HOOK using the benefit-led suspense formula (Part 4a).",
    "Write the full clean narration script using the retention structure (Part 4b), to the target length (Part 4c).",
    "Compute total duration (words ÷ 2.5) and segment into 3–5s FRAMES (~8–13 words each); assign timestamps (Part 5).",
    "Build the shot list table (Deliverable A): for each FRAME write the narration slice, a one-line visual, the "
    "element breakdown + draw order, prompt #s, and type.",
    "Build the PROMPT PACK (Deliverable B, Part 8f): a Consistency Header (compact style tag + character/prop bible + "
    "settings) plus per-element prompts grouped by frame, each a single isolated element on a transparent background "
    "and restating recurring characters VERBATIM so all elements stay consistent.",
    "Generate the ELEMENTS (Part 8g): run each prompt on Seedream 4.5 (free in the web app), reuse Host/prop Elements "
    "+ style key-frame, then REMOVE BACKGROUND → save transparent PNGs into frames/frame_0NN/. Nano Banana Pro only "
    "for text elements.",
    "For any optional CLIP shots, generate the element still then animate it (Part 8c).",
    "Assemble in VideoScribe (Part 8h): the hand draws each element on, frame by frame; add VO/captions/music; "
    "render MP4.",
    "Save Deliverable A (shot-list.docx), Deliverable B (prompt-pack.txt), and Deliverable C (frames/ folder, + "
    "videoscribe/ and clips/ if any) per Part 9 naming.",
    "Run the checklists (Part 11) before delivering.",
]:
    numbered(doc, step)

# ============================ 11. CHECKLISTS ============================
h1(doc, "11. Quality Checklists")
h2(doc, "Script")
for c in [
    "Hook is benefit-led and suspenseful (a question/what-if), and does NOT just name the concept.",
    "Body follows: stakes → animal → parallel → where-it-breaks-down → payoff → button.",
    "A question or open loop roughly every 20–30 seconds.",
    "Reads naturally aloud; short concrete sentences; matches target runtime.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2(doc, "Shot list & timing")
for c in [
    "Every frame is 3–5 seconds; each frame lists its elements + draw order.",
    "Timestamps are cumulative and sum to total duration.",
    "Each frame has narration, visual, element breakdown + draw order, prompt #s, and type.",
    "Numbers match across shot list, prompt pack, and frame/element files.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2(doc, "Prompt pack (Deliverable B)")
for c in [
    "A Consistency Header is present: compact style tag + character/prop bible + palette + settings.",
    "Every per-element prompt restates recurring characters VERBATIM (never 'same as above').",
    "Every prompt asks for a single isolated element on a transparent background and is under 3000 characters.",
    "Grouped by frame with draw order; numbered to match; seed noted.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2(doc, "Element PNGs (each)")
for c in [
    "Brand palette only; tomato coral reserved for the one focal element of the frame.",
    "Bright clip-art look: soft shading + gentle highlights, thin edge accents; no heavy black outlines, no grain, "
    "no gradients, no 3D.",
    "Single isolated element, TRANSPARENT background, clean cutout (no halo), no ground shadow.",
    "Recurring characters identical to their canonical Element; named frame_0NN_elXX_<name>.png.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2(doc, "Assembly")
for c in [
    "16:9 cream canvas; clear focal subject per frame; bottom 15% caption-safe.",
    "Elements drawn on in the planned order; timing matches the VO; frames flow in sequence.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")

doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run("End of guide. Inputs: finance_animal_behavior_matches.xlsx. Outputs per row: a shot-list .docx, a "
                 "Higgsfield prompt pack, and a frames folder of transparent element PNGs — assembled in VideoScribe "
                 "with the animated hand-reveal, in the fixed series style, with visuals changing every 3–5 seconds.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string("888888")

out = "/home/user/dataexploration/PRODUCTION_GUIDE.docx"
doc.save(out)
print("Saved:", out)
