#!/usr/bin/env python3
"""Build STYLE_AND_COLOR_GUIDE.docx — the single source of truth for color-palette
and visual-style consistency across EVERY video in the series. Self-contained and
copy-paste-ready: the one document to keep open / upload so all episodes match."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- canonical blocks (must match the master guide) ----------------
PALETTE = [
    ("Sunshine Yellow", "#FBC02D", "Beaks, manes, warm highlights"),
    ("Sky Blue", "#4FC3E8", "Water/sky elements, calm characters, the Host"),
    ("Leaf Green", "#7CC576", "Foliage, ground shapes"),
    ("Tomato Coral", "#F0654E", "THE focal color — the ONE key idea per frame (sparingly)"),
    ("Tangerine", "#F39237", "Secondary warm accent (lion/monkey)"),
    ("Blush Pink", "#F6A9B6", "Cheek blush, soft accents"),
    ("Stone Gray", "#9DB0BC", "Elephant/rhino bodies, neutral props"),
    ("Warm Brown", "#8D6E63", "Trunks, wood, branches"),
    ("Cream", "#FBF7EF", "The VideoScribe canvas / negative space"),
    ("Charcoal", "#3A3631", "Eyes, edge accents, fine detail — never pure black"),
]

FULL_SUFFIX = (
    "Bright, cheerful, wholesome flat 2D vector children's-book clip-art illustration. Rounded, cute, friendly "
    "characters and objects with clean simple shapes and chunky proportions. Soft subtle shading and gentle soft "
    "highlights for a little depth; thin minimal darker edge accents only where a shape needs separation (no heavy "
    "black outlines). Small dot eyes, rosy cheek blush, simple happy expressions. Bright multicolor brand palette: "
    "sunshine yellow (#FBC02D), sky blue (#4FC3E8), leaf green (#7CC576), tomato coral (#F0654E), tangerine "
    "(#F39237), blush pink (#F6A9B6), stone gray (#9DB0BC), warm brown (#8D6E63), cream (#FBF7EF), charcoal "
    "(#3A3631). SINGLE isolated element only, centered, on a plain TRANSPARENT background — no scene, no setting, no "
    "ground shadow. Clean, modern, wholesome. No grain, no gradients, no 3D, no realistic textures, no photorealism, "
    "no heavy black outlines, no clutter, no scene background."
)
COMPACT_SUFFIX = (
    "STYLE: bright cheerful flat 2D vector children's-book clip-art; rounded cute friendly shapes, soft subtle "
    "shading + gentle highlights, thin minimal edge accents (no heavy black outlines); dot eyes, rosy cheeks; "
    "palette sunshine #FBC02D, sky #4FC3E8, leaf #7CC576, coral #F0654E, tangerine #F39237, blush #F6A9B6, stone "
    "#9DB0BC, brown #8D6E63, cream #FBF7EF, charcoal #3A3631; SINGLE isolated element, centered, TRANSPARENT "
    "background, no scene/no ground shadow; clean modern wholesome. No grain/gradient/3D/realistic/photoreal/heavy "
    "outline/clutter/scene-background."
)
NEGATIVE = (
    "scene background, environment, ground shadow, multiple subjects, photorealism, 3D render, CGI, realistic "
    "textures, film grain, gradients banding, glossy/metallic highlights, ambient occlusion, drop shadows, lens "
    "flare, glow/bloom, heavy black outlines, comic linework, sketchy wobble, crosshatching, dark/gritty mood, "
    "cluttered detail, complex detailed faces, tiny illegible text, watermarks, extra limbs/fingers, natural "
    "photographic animal coloring."
)

HOST = "cute clip-art fox host (rounded chunky sky-blue body, charcoal dot eyes, rosy cheeks, tomato-coral knitted scarf)"
BAT = "the 'animal of the week' as a cute clip-art character, re-tinted into the brand palette (NOT natural photo colors), rounded and friendly, dot eyes, rosy cheeks"
FIG = "simple rounded clip-art person figure (sky-blue body, charcoal dot eyes, no other facial detail); a crowd is many identical figures, the odd-one-out is tomato coral"

# ---------------- helpers ----------------
def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexv.lstrip("#"))
    tcPr.append(shd)


def sct(cell, text, bold=False, size=9, color=None):
    cell.text = ""; r = cell.paragraphs[0].add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def h1(t): return doc.add_heading(t, level=1)


def body(t, size=10.5, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(size); r.italic = italic
    return p


def bullet(t, size=10.5):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t).font.size = Pt(size); return p


def box(label, text):
    if label:
        lp = doc.add_paragraph(); r = lp.add_run(label); r.bold = True
        r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string("1F4E5F")
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); shade(c, "F2F4F0"); c.text = ""
    rr = c.paragraphs[0].add_run(text); rr.font.name = "Consolas"; rr.font.size = Pt(9)
    doc.add_paragraph()


def callout(text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); shade(c, "FBEFD6"); c.text = ""
    r = c.paragraphs[0].add_run(text); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("8A4B16"); doc.add_paragraph()


doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

# ---------------- title ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("STYLE & COLOR CONSISTENCY GUIDE"); r.bold = True; r.font.size = Pt(26)
r.font.color.rgb = RGBColor.from_string("1F4E5F")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Finance × Animal Behavior — the ONE look every video must share")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string("F0654E")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Keep this open while generating. Every image in every episode obeys this page. "
              "Only the CONTENT changes between videos — the palette and style never do.")
r.italic = True; r.font.size = Pt(10)
doc.add_paragraph()

callout("THE ONE RULE: same palette + same style on every frame of every episode. That sameness IS the brand. "
        "If a new image doesn't match this page, regenerate it.")

# ---------------- 1. THE STYLE ----------------
h1("1. The Style — in one sentence")
body("Bright, cheerful, wholesome FLAT 2D vector children's-book CLIP-ART: rounded, cute, friendly characters and "
     "objects; clean simple shapes; soft subtle shading with gentle highlights; thin minimal edge accents (no heavy "
     "black outlines); small dot eyes and rosy cheeks. Every asset is a SINGLE isolated element on a transparent "
     "background, so it can be drawn onto the cream canvas in VideoScribe.")
h1("2. Style rules (the non-negotiables)")
for x in [
    "CLIP-ART, ISOLATED — one element per asset, centered, on a TRANSPARENT background. No scenes, no settings, no ground shadow.",
    "ROUNDED & CUTE — chunky friendly proportions; circles, ovals, soft rounded shapes.",
    "SOFT SHADING + HIGHLIGHTS — gentle subtle depth, never flat-dead and never 3D.",
    "THIN EDGE ACCENTS ONLY — light darker edges where shapes need separation; NO heavy black ink outlines.",
    "BRAND PALETTE ONLY (Section 3) — bright and cheerful, but fixed. Re-tint real animals into the palette.",
    "FRIENDLY FACES — small dot eyes, rosy cheek blush, simple happy expressions.",
    "CLEAN & WHOLESOME — uncluttered, modern; NO grain, NO texture overlays, NO gradients.",
    "ONE FOCAL COLOR — reserve TOMATO CORAL for the single thing the viewer should look at in each frame.",
]:
    bullet(x)

# ---------------- 3. PALETTE ----------------
h1("3. The Color Palette (canonical — use these exact hex values)")
pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
for i, x in enumerate(["Swatch", "Name", "Hex", "Where it's used"]):
    sct(pt.rows[0].cells[i], x, bold=True, size=9, color="FFFFFF"); shade(pt.rows[0].cells[i], "1F4E5F")
for name, hexv, usage in PALETTE:
    c = pt.add_row().cells
    shade(c[0], hexv); sct(c[0], "", size=9)
    sct(c[1], name, bold=True, size=9); sct(c[2], hexv, size=9); sct(c[3], usage, size=9)
for row in pt.rows:
    row.cells[0].width = Inches(0.7); row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(0.9); row.cells[3].width = Inches(3.4)
doc.add_paragraph()
bullet("Bright but balanced: spread the palette across elements; let TOMATO CORAL spotlight the one key idea.")
bullet("Cream (#FBF7EF) is the canvas/background. No pure black and no pure white — use Charcoal and Cream.")
bullet("Tints/shades of these colors are fine; brand-new colors are not.")

# ---------------- 4. CHARACTER ANCHORS ----------------
h1("4. Character anchors (describe with the SAME words every time)")
body("Re-stating these exact descriptions in every prompt is what keeps recurring characters identical across all "
     "videos. Save the Host (and any repeated prop) as a Higgsfield ELEMENT and reuse it — don't redraw it.")
box("HOST (recurring narrator — in every episode):", HOST)
box("ANIMAL OF THE WEEK (the episode's animal):", BAT)
box("CROWDS / MARKETS / 'YOU':", FIG)

# ---------------- 5. THE STYLE TAGS ----------------
h1("5. The style tags (paste at the end of EVERY prompt)")
body("Append one of these verbatim to every image prompt so the look is locked. Use the COMPACT tag for the "
     "self-contained per-element prompts; use the FULL suffix for one-off hero elements where length isn't a concern.")
box("COMPACT STYLE TAG (use for normal per-element prompts):", COMPACT_SUFFIX)
box("FULL MASTER STYLE SUFFIX (use for single hero elements):", FULL_SUFFIX)
box("NEGATIVE / AVOID (steer away from all of these):", NEGATIVE)

# ---------------- 6. CONSISTENCY CHECKLIST ----------------
h1("6. Per-image consistency checklist")
for c in [
    "Brand palette ONLY; tomato coral reserved for the one focal element.",
    "Clip-art look: soft shading + gentle highlights, thin edge accents — no heavy black outlines, no grain, no gradients, no 3D.",
    "SINGLE isolated element on a clean TRANSPARENT background — no scene, no ground shadow, no halo from cutout.",
    "Recurring characters described with the exact anchor words / generated from the saved Element.",
    "Dot eyes + rosy cheeks; rounded, cute, friendly proportions.",
    "Matches every other image in the episode — and every other episode.",
]:
    doc.add_paragraph("☐  " + c, style="List Bullet")

note = doc.add_paragraph()
r = note.add_run("This is the consistency source of truth. It is the condensed art-direction core of the "
                 "PROJECT MASTER GUIDE; if they ever disagree, the master guide wins and this page should be regenerated.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string("888888")

out = "/home/user/dataexploration/STYLE_AND_COLOR_GUIDE.docx"
doc.save(out)
print("Saved:", out)
