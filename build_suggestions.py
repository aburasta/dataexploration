from docx import Document
from docx.shared import Pt, RGBColor

doc = Document()

def h1(t):
    p = doc.add_heading(t, level=1)
def h2(t):
    p = doc.add_heading(t, level=2)
def h3(t):
    p = doc.add_heading(t, level=3)
def para(t, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p
def label(t, color):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.color.rgb = color
    return p

RED = RGBColor(0xB0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x66, 0x00)
BLUE = RGBColor(0x1F, 0x3A, 0x93)

doc.add_heading('Revision Suggestions for Research Proposal', 0)
para('“Impact of Leadership on Cost Control in Mining Operations in Ghana: A Case Study of Chirano Mine” — Bright Okrah (10332534)', italic=True)
para('This document walks through every paragraph the reviewer flagged. For each one it shows the section heading, the ORIGINAL paragraph as written in the proposal, and a SUGGESTED REWRITE that responds to the review comments. Original text is shown in red, suggested rewrites in green.', italic=True)

# --- Section 1.3 ---
h1('1.3 Problem Statement')

h3('Paragraph 1 (opening of Section 1.3)')
label('ORIGINAL:', RED)
para('Mining companies in Ghana operate under steady cost pressure because of increase in input cost such as labour cost, fuel cost, materials and maintenance cost, dependence on contractors, exchange rate pressures and inefficiencies push operating costs up. The significant increase in gold prices on the international market in recent years has led to suppliers increasing prices because they know miners can “afford” higher costs, labour demanding increase in wages and benefits and contractors seeking upward review of contract rates. This situation has brought cost pressure challenges to the leadership of mining companies in Ghana and as such cost control has become a major key performance indicator in assessing the performance of the managers of the mine. In response, mines put in place cost-control systems but still there are budget overruns and uneven cost performance (Drury, 2018).', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Mining companies in Ghana operate under steady cost pressure driven by rising labour, fuel, materials and maintenance costs, heavy reliance on contractors, exchange-rate volatility and operational inefficiencies. Recent increases in the international gold price have compounded this pressure, as suppliers raise prices, labour presses for higher wages and benefits, and contractors seek upward reviews of their rates. While these drivers are largely external to any single mine, the internal managerial response to them is not: leadership determines whether formal cost-control systems are enforced, whether departments coordinate, and whether budget variances trigger corrective action. It is for this reason that cost control has become a key performance indicator for mine managers in Ghana, and yet, despite the installation of cost-control systems, budget overruns and uneven cost performance persist (Drury, 2018). This points less to the absence of controls than to weakness in the managerial lever that makes them bite.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes that the paragraph names external drivers of cost pressure and then pivots to leadership without bridging the two. The rewrite acknowledges the external drivers first, then explicitly positions leadership as the controllable internal lever, which sets up the study’s focus logically instead of by assertion.', color=BLUE)

h3('Paragraph 2 (mid-Section 1.3)')
label('ORIGINAL:', RED)
para('One plausible explanation is that formal cost systems due to leadership and coordination challenges or lack of corrective action and little accountability in cases of non-performance (Otley, 1999; Northouse, 2022).', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('One plausible explanation is that formal cost systems fail to deliver their intended results because of leadership and coordination weaknesses, the absence of timely corrective action, and limited accountability when cost targets are missed (Otley, 1999; Northouse, 2022).', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer explicitly flagged this as a grammatically incomplete sentence. The rewrite completes the verb, keeps the same argument, and preserves the citations.', color=BLUE)

h3('Paragraph 3 (closing of Section 1.3)')
label('ORIGINAL:', RED)
para('This naturally implies that even though Chirano Mine have put in place some cost measures (Asante Gold Corporation, n.d.), there is a little structured evidence on how leadership styles and managerial practices actually influence cost outcomes in this setting. The problem this study addresses, therefore, is the limited empirical understanding of how leadership affects cost control in mining operations in Ghana which means cost actions may be overlooking the leadership conditions needed to make those controls work.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Although Chirano Mine has put in place a number of cost measures (Asante Gold Corporation, n.d.), and although prior work has linked leadership to organisational performance in Ghana (Donkor, Dongmei & Sekyere, 2021) and management control to cost outcomes more generally (Otley, 1999; Drury, 2018), no study to date examines the specific mechanisms by which leadership styles and everyday managerial practices shape cost-control performance in a Ghanaian large-scale mine. The problem this study addresses is therefore the limited empirical understanding of how leadership influences cost-control outcomes at Chirano Mine, with the consequence that cost interventions in this setting may continue to overlook the leadership conditions on which those controls depend.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes two problems: the gap is asserted rather than drawn from evidence, and its scope (“mining operations in Ghana”) does not match the single-case design. The rewrite cites the actual literature to justify the gap, and narrows the scope to Chirano so it aligns with the case study.', color=BLUE)

# --- Section 1.6.1 ---
h1('1.6.1 Research Hypotheses')

h3('Hypothesis paragraph')
label('ORIGINAL:', RED)
para('Because the study includes a quantitative component, the following hypotheses will be tested:\nH0: Leadership has no significant effect on cost-control performance in mining operations.\nH1: Leadership has a significant effect on cost-control performance in mining operations.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Drawing on the Full-Range Leadership Model (Bass & Riggio, 2006) and the empirical evidence reviewed in Section 1.8.2, the study tests the following directional, style-specific hypotheses:', color=GREEN)
para('H1a: Transformational leadership has a positive and significant effect on cost-control performance at Chirano Mine.', color=GREEN)
para('H1b: Transactional leadership (contingent reward and active management-by-exception) has a positive and significant effect on cost-control performance at Chirano Mine.', color=GREEN)
para('H1c: Passive-avoidant leadership has a negative and significant effect on cost-control performance at Chirano Mine.', color=GREEN)
para('Objectives 1, 2 and 4 are descriptive and are answered without inferential testing; only Objective 3 (the relationship between leadership and cost-control performance) is tested through the hypotheses above.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer criticises the single blanket, non-directional hypothesis for wasting the multi-theory framework the proposal has built and for being placed before the literature review. This rewrite (i) makes the hypotheses style-specific, (ii) makes them directional as the theory predicts, (iii) makes clear which objectives are descriptive versus inferential, and (iv) signals that the hypotheses derive from the literature review. The block itself should also be MOVED to the end of Section 1.8 (after the empirical review), not kept in Section 1.6.1.', color=BLUE)

# --- Section 1.8.1 ---
h1('1.8.1 Theoretical Literature Review')

h3('Paragraph 1 (Transformational leadership)')
label('ORIGINAL:', RED)
para('The first theoretical underpinning of this study is transformational leadership, which grew out of Burns (1978) which initially highlighted the difference between leaders who transform their followers and those who simply exchange rewards for effort. Bass and Riggio (2006) developed this idea further, describing transformational leaders as those who set a clear vision, encourage fresh thinking and pay attention to individual subordinates. In a cost-control setting, transformational leadership may help build a culture of cost culture throughout the entire organization cost.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('The first theoretical underpinning of this study is transformational leadership. Burns (1978) originally distinguished leaders who transform their followers from those who merely exchange rewards for effort, and Bass and Riggio (2006) developed the construct into four measurable components: idealised influence, inspirational motivation, intellectual stimulation and individualised consideration. In a cost-control setting these components matter because they shape whether staff internalise cost discipline as a shared value, raise cost problems early, and propose efficiency improvements rather than treat cost targets as an externally imposed constraint. The four components will provide the survey constructs used to measure transformational leadership in this study.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer flags the garbled closing sentence (“a culture of cost culture throughout the entire organization cost”) and, more importantly, notes that the theory is not developed into the measurable dimensions the study will need. The rewrite names the four components explicitly, links them to cost control, and signals that they become the survey items.', color=BLUE)

h3('Paragraph 2 (Transactional leadership)')
label('ORIGINAL:', RED)
para('The second is transactional leadership, which Burns (1978) presented as the counterpart to the transforming style and which Bass and Riggio (2006) also elaborated on. This style  of leadership is closely tied to cost control because the actual mechanisms put in place to check costs may rely or work on accountability all rest on transactional practices that tie employee rewards or punishment to achievement of cost targets.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('The second theory is transactional leadership, which Burns (1978) presented as the counterpart to the transforming style and which Bass and Riggio (2006) elaborated through two active dimensions — contingent reward and active management-by-exception — as well as the passive-avoidant dimensions of passive management-by-exception and laissez-faire behaviour. Transactional leadership sits at the heart of cost control because its accountability mechanisms are precisely what formal cost systems rely on: contingent reward ties recognition and bonuses to the achievement of cost targets, while active management-by-exception ensures that variances are identified and corrected before they escalate. These same dimensions will be used as survey constructs, allowing the study to test whether it is the active transactional behaviours in particular that carry cost-control performance.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes that transactional leadership is described “without contingent reward and management-by-exception,” which are the dimensions most directly tied to cost accountability. The rewrite names those dimensions, links each to a cost-control mechanism, and fixes the awkward phrasing.', color=BLUE)

h3('Paragraphs 3 & 4 (Situational leadership and contingency theory)')
label('ORIGINAL:', RED)
para('The third theory is situational leadership, associated with Hersey, Blanchard and Johnson (2001), which holds that there is no single best way to lead. Effective leaders adjust their behaviour to the competence and commitment of their staff and to the demands of the moment. Because of the complexity of mining operations, leaders may need to shift their approach constantly to adjust and deal with rising needs for costs.', color=RED)
para('The fourth is contingency theory (Fiedler, 1967), which similarly argues that the effectiveness of a leadership style depends on the situation, including the structure of the task, the relationship between leader and team and the leader’s influence over the team in general to achieve the set goals. Taken together, these four theories suggest that cost outcomes are likely to relate to how well leaders apply different styles as conditions change.', color=RED)
label('SUGGESTED REWRITE (combine into one shorter paragraph as secondary framing):', GREEN)
para('Because situational leadership (Hersey, Blanchard & Johnson, 2001) and contingency theory (Fiedler, 1967) make substantially the same claim — that no single leadership style is universally effective and that its impact depends on task structure, team competence and leader–member relations — they are treated together in this study as secondary framing rather than as separate primary theories. Their contribution here is to explain why the effect of transformational and transactional leadership on cost control at Chirano Mine may itself vary across departments and operational contexts (for example, between open-pit and underground operations, or between finance and operations teams). The primary theoretical anchor of the study therefore remains the Full-Range Leadership Model, with situational-contingency reasoning used to interpret contextual variation in the results.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer’s core theoretical concern is breadth: “situational and contingency theory make the same basic point … so four theories risk a scattered framework in which none is developed fully,” and recommends anchoring the study in the Full-Range Leadership Model with the other two as secondary framing. The rewrite collapses the two overlapping theories into a single secondary paragraph, states this choice explicitly, and repairs the garbled clause (“rising needs for costs”).', color=BLUE)

# --- Section 1.8.2 ---
h1('1.8.2 Empirical Literature Review')

h3('Full section (currently one short paragraph, expand to a proper review)')
label('ORIGINAL:', RED)
para('The wider empirical literature supports the view that leadership shapes organisational performance. In a large meta-analysis covering many studies and settings, Judge and Piccolo (2004) consistently found that transformational and transactional leadership was strongly and positively related to a range of performance criteria.\nIn the Ghanaian context, Donkor, Dongmei and Sekyere (2021) studying state-owned enterprises in Ghana, found that transformational leadership had a positive and significant relationship with employee performance. On the cost side, Drury (2018) describes how management review functions as the core technique for planning and controlling expenditure. Otley (1999) makes a related point that performance and control should be understood as a whole rather than something to be treated in isolation. The studies above point strongly to a relationship, but the specific pathways and the strength of the link in a mining environment are still under-documented and this study is designed to help close that gap.', color=RED)
label('SUGGESTED REWRITE (structure and depth to aim for):', GREEN)
para('The empirical review should be reorganised into three thematic strands rather than a single list, and expanded so each strand contains at least three to five studies whose designs, samples and findings are compared, not merely cited. A suggested structure:', color=GREEN)
para('(a) Leadership and general organisational performance. Anchor with Judge and Piccolo (2004) as the meta-analytic benchmark, but add subsequent syntheses (for example Wang, Oh, Courtright & Colbert, 2011) and appraise their designs so the reader sees the strength and limits of the leadership–performance link, not just its direction.', color=GREEN)
para('(b) Leadership and financial/cost performance. This is the strand the current review is missing. Locate studies that specifically link leadership styles (transformational, transactional, and passive-avoidant) to cost efficiency, budget adherence, or financial performance in industrial, manufacturing or extractive settings. Where mining-specific studies are scarce, complement them with studies from manufacturing and construction, which share the operational cost profile. For each study, report the setting, sample, measures of leadership and of cost performance, and the direction and size of the effect.', color=GREEN)
para('(c) Leadership and performance in Ghana / West Africa and in the mining sector. Retain Donkor, Dongmei and Sekyere (2021), and add studies of leadership in Ghanaian firms and of management control and cost performance in African mining. This strand localises the evidence and shows what has and has not been examined at home.', color=GREEN)
para('The section should then synthesise across the three strands rather than list them: which leadership styles are most consistently linked to cost outcomes, under what conditions, and with what methodological limitations. It should close on a gap that is earned from the review — for example, that although transformational and transactional leadership are consistently linked to general performance, and although management control is well developed in the cost literature, few studies have directly tested style-specific effects of leadership on cost-control performance in a Ghanaian large-scale mine — and it is at this point that the style-specific hypotheses (H1a–H1c) should be stated.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer calls this “the weakest section” because it is a single short paragraph, cites work on leadership-and-performance and on cost accounting in general but nothing at the leadership-and-cost-control nexus, and lists rather than synthesises. The rewrite reorganises the review thematically, forces coverage of the missing intersection, sets an expectation for synthesis and methodological appraisal, and moves the hypothesis specification to the end of the review as the reviewer requested.', color=BLUE)

# --- Section 1.8.3 ---
h1('1.8.3 Definition of Terms and Conceptual Framework')

h3('Paragraph 2 (framework paragraph)')
label('ORIGINAL:', RED)
para('The study treats leadership as the independent variable and cost-control performance as the dependent variable. Leadership is measured through the dimensions named in the theoretical review and cost-control performance is reflected in indicators that are relevant to how Chirano Mines is currently tracking their costs as performance indicators. The central proposition is that stronger leadership practices improve operational efficiency in controlling cost. A diagram of this framework will be presented in the full study.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('The study treats leadership as the independent variable, operationalised through the four transformational components (idealised influence, inspirational motivation, intellectual stimulation, individualised consideration) and the transactional dimensions of contingent reward and active management-by-exception, following the Full-Range Leadership Model. Cost-control performance is the dependent variable and is operationalised through four indicators drawn from Chirano Mine’s existing performance-management practice: (i) adherence to approved operating budgets (variance from budget), (ii) unit operating cost per ounce, (iii) timeliness of corrective action on cost variances, and (iv) achievement of departmental cost-reduction targets. Accountability and coordination, discussed in the background as the everyday mechanisms through which leadership acts on cost outcomes, are modelled as mediating variables that carry the effect of leadership styles onto cost-control performance. The central proposition is that stronger transformational and active transactional leadership improves cost-control performance, mediated by stronger accountability and coordination. The conceptual framework is shown in Figure 1 [insert diagram], with arrows from each leadership dimension through the two mediators to the four cost-control indicators.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes that the four theories are not linked to specific variables, that cost-control performance is left vague, that the framework is deferred to the full study rather than shown, and that accountability and coordination — mentioned in the background — read like mediators but are not modelled. The rewrite names the leadership dimensions, gives cost-control performance four concrete indicators, promotes accountability and coordination to explicit mediators, and commits to including the diagram now rather than later.', color=BLUE)

# --- Section 1.9.1 ---
h1('1.9.1 Research Approach')

h3('Approach paragraph')
label('ORIGINAL:', RED)
para('The study will adopt a qualitative approach for confirming empirically the connection between the leadership styles being implemented at Chirano Mines and the performance of their cost control initiatives. The quantitative approach will allow for the relationship to be tested adequately so the definitive relationship and size of the influence can be measured.', color=RED)
label('SUGGESTED REWRITE (option A — pure quantitative, recommended if interviews are dropped):', GREEN)
para('The study will adopt a quantitative research approach. A quantitative approach is appropriate because the central objective is to test the relationship between leadership styles and cost-control performance at Chirano Mine and to estimate the size of that relationship, which requires structured measurement and inferential analysis. Data will be gathered through a structured Likert-scale questionnaire and analysed using descriptive statistics, correlation and multiple regression, in line with the style-specific hypotheses set out at the end of Section 1.8.', color=GREEN)
label('SUGGESTED REWRITE (option B — explanatory mixed-methods, if interviews are retained):', GREEN)
para('The study will adopt an explanatory sequential mixed-methods approach. The dominant strand is quantitative: a structured Likert-scale questionnaire will be used to test the style-specific hypotheses through correlation and multiple regression. A subsequent qualitative strand will consist of a small number of semi-structured interviews with finance, operations and management staff, whose purpose is to explain and contextualise the quantitative findings — in particular, how leadership behaviours actually operate through accountability and coordination to shape cost outcomes at Chirano Mine.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer identifies the paragraph as internally contradictory (it labels itself qualitative but describes a quantitative design) and notes that interviews appear later without being introduced. The two options resolve the contradiction: either commit fully to quantitative, or acknowledge the interviews and name the design explanatory mixed-methods. Either is defensible; the student must pick one and be consistent through the rest of the methodology.', color=BLUE)

# --- Section 1.9.2 ---
h1('1.9.2 Research Design')

h3('Design paragraph')
label('ORIGINAL:', RED)
para('A case study design is being used here with Chirano Mine as the single case. The case study is appropriate because of the lack of available evidence in this specific inquiry across the Ghanaian mining industry which means, the use of a case study contextualizes the findings adequately enough for further use and improvement to other areas of the mining sector.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('The study uses a single-organisation cross-sectional survey design, with Chirano Mine as the study site. Although the study is often described colloquially as a “case study” because it focuses on a single organisation, its analytical logic is that of a cross-sectional survey: primary data will be collected at one point in time from a sample of staff within one organisation and analysed to test hypothesised relationships between leadership styles and cost-control performance. A single-organisation design is appropriate because it allows deep access to the cost-control practices, documentation and staff of one large-scale Ghanaian mine — access that would be difficult to obtain across multiple sites — while still supporting statistical testing of the study’s hypotheses.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes that a single-case label is loose for what is actually a hypothesis-testing survey within one organisation, and recommends renaming the design as a single-organisation cross-sectional survey. The rewrite adopts that language and justifies the choice.', color=BLUE)

# --- Section 1.9.4 ---
h1('1.9.4 Sampling')

h3('Sampling paragraph')
label('ORIGINAL:', RED)
para('For the study, a sample of about 80 to 120 respondents will be targeted, subject to access and approval. Stratified random sampling will be used so that the main departments are properly represented in the survey.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Chirano Mine employs approximately [insert total staff figure] permanent employees, of whom an estimated [insert figure] are located in the finance, operations, processing and management functions that form the study’s target population. A stratified random sampling strategy will be used, with the four functions treated as strata and respondents allocated proportionally to the size of each stratum. Sample size is determined by power analysis for the regression model: following Green’s (1991) rule of N ≥ 50 + 8k for testing the overall model, with k = 6 leadership predictors (four transformational components plus contingent reward and active management-by-exception), the minimum required sample is 98 respondents. G*Power was used to confirm that a sample of 100 provides power of 0.80 to detect a medium effect (f² = 0.15) at α = 0.05 for six predictors. A target sample of 110 respondents will therefore be pursued to allow for non-response, allocated across the four strata as follows: [insert proportional allocation table].', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer flags that “80 to 120 respondents, subject to access” is stated without a determination formula and that strata and allocations are not specified. The rewrite justifies the sample size with a named formula and a G*Power calculation matched to the regression design, and promises an explicit proportional allocation. The bracketed figures should be replaced with the actual staff numbers once available from the company.', color=BLUE)

# --- Section 1.9.5 ---
h1('1.9.5 Instrumentation')

h3('Instrumentation paragraph')
label('ORIGINAL:', RED)
para('The data instrument to be used will be a structured questionnaire made up mainly of closed-ended items measured on a five-point Likert scale. It will gather information on leadership style and perceptions of cost-control performance as well as  pre-formatted spreadsheets to collect the cost related information. The questionnaire items will be adapted from established leadership and management-control instruments used in this study for the theoretical foundations. The instruments will be checked for validity through supervisory review and the internal consistency using Cronbach’s alpha.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Data will be collected through two instruments. First, a structured questionnaire made up of closed-ended items measured on a five-point Likert scale. Leadership will be measured using items adapted from the Multifactor Leadership Questionnaire (MLQ-5X; Bass & Avolio, 2004), retaining the four transformational subscales (idealised influence, inspirational motivation, intellectual stimulation, individualised consideration) and the active transactional subscales (contingent reward, active management-by-exception). Cost-control performance will be measured through two complementary sources: (i) a perceptual scale in the same questionnaire, with items covering budget adherence, timeliness of corrective action on variances, and achievement of departmental cost-reduction targets, adapted from Otley (1999) and prior management-control instruments; and (ii) objective cost data extracted from Chirano Mine’s management reports over the most recent 24 months, covering unit operating cost per ounce, actual-versus-budget variance and cost-reduction target attainment at departmental level. The two sources will be triangulated at the departmental level so that perceived cost-control performance can be checked against reported cost outcomes. Instrument validity will be established through supervisory and expert review and by piloting; internal-consistency reliability will be assessed with Cronbach’s alpha (α ≥ 0.70 accepted as satisfactory for each construct).', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer flags that leadership is not broken into named dimensions with a source instrument, that cost-control performance has no operational definition (the most serious omission), and that the mix of survey data with pre-formatted cost spreadsheets is unexplained. The rewrite names the MLQ-5X as the source instrument, lists the subscales, gives cost-control performance both a perceptual scale and an objective indicator set, and explains how the two data sources are combined.', color=BLUE)

# --- Section 1.9.7 ---
h1('1.9.7 Data Analysis')

h3('Data analysis paragraph')
label('ORIGINAL:', RED)
para('The data collected will be coded and analysed using descriptive statistics to initially provide the context within which further inquiry will be sought and final interpretations will be done. The inferential analysis to be used includes correlation and regression in order to test the relationship between the leadership variables and cost-control performance and to evaluate the hypothesis of the study.', color=RED)
label('SUGGESTED REWRITE:', GREEN)
para('Analysis will proceed in three stages, each mapped to specific research questions. First, descriptive statistics (frequencies, means and standard deviations) will be used to profile the leadership styles practised at Chirano Mine (RQ1) and the cost-control mechanisms currently in use (RQ2), and to identify the leadership-related challenges most frequently reported by staff (RQ4). Second, Pearson correlation and multiple linear regression will be used to test the relationship between the six leadership predictors and cost-control performance (RQ3, H1a–H1c). The regression model will be Cost-Control Performance = β0 + β1·IdealisedInfluence + β2·InspirationalMotivation + β3·IntellectualStimulation + β4·IndividualisedConsideration + β5·ContingentReward + β6·ActiveMBE + ε, with departmental controls entered as covariates. Standard regression assumptions (linearity, normality of residuals, homoscedasticity, and multicollinearity via VIF) will be checked before interpretation. Third, where mediation by accountability and coordination is proposed by the conceptual framework, a bootstrapped mediation analysis (Hayes PROCESS) will be run. If semi-structured interviews are retained under the mixed-methods design, they will be analysed with thematic analysis following Braun and Clarke (2006) to explain and contextualise the quantitative results. Company cost data covering the most recent 24 months of monthly management reports will be obtained under a data-sharing arrangement with the mine, with all firm-identifying detail treated as confidential and reported only in aggregated form.', color=GREEN)
label('WHY THIS CHANGE:', BLUE)
para('The reviewer notes that the analysis plan covers only RQ3 and says nothing about RQ1, RQ2, RQ4 or the interviews, that the regression model and its assumptions are not specified, and that the company cost data and its period and confidentiality terms are not stated. The rewrite maps each stage of analysis to specific research questions, writes out the regression model and the assumption checks, adds a mediation analysis consistent with the revised conceptual framework, provides for thematic analysis of interviews if the mixed-methods option is chosen, and specifies the secondary cost data and its confidentiality terms.', color=BLUE)

# --- Overall priority list ---
h1('Priority order for implementing these revisions')
para('If time is limited, the reviewer’s priority order is: (1) resolve the qualitative-versus-quantitative contradiction in Section 1.9.1 and rename the design in 1.9.2; (2) replace the single hypothesis with the style-specific H1a–H1c and move them to the end of Section 1.8; (3) operationally define cost-control performance in Sections 1.8.3 and 1.9.5; (4) justify the sample size in Section 1.9.4; (5) expand and synthesise the empirical review in Section 1.8.2; (6) draw the conceptual framework and align the gap with the single-case scope in Sections 1.3 and 1.8.3; and (7) proofread the garbled sentences in Sections 1.3 and 1.8.1.', italic=True)

out = '/home/user/dataexploration/Revision_Suggestions_Bright_Proposal.docx'
doc.save(out)
print('Saved:', out)
