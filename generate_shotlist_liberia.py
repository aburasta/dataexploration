#!/usr/bin/env python3
"""
Build the TIMED SHOT-LIST Word doc for the Liberia 1985 (Quiwonkpa) voiceover.
Footage library assembled directly from reputable sources (footage sub-agents
hit a session limit), so links are a starting map to VERIFY, not cleared rights.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches

WPM = 140.0
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TC = RGBColor(0x8A, 0x1C, 0x1C)
TYPE_COLOR = {
    "REAL": RGBColor(0x2E, 0x7D, 0x32), "CC": RGBColor(0x1F, 0x6F, 0xEB),
    "ARCHIVE": RGBColor(0xB7, 0x6E, 0x00), "STOCK": RGBColor(0x66, 0x66, 0x66),
    "DOC": RGBColor(0x6A, 0x1B, 0x9A),
}

A = {
    "crowds85": ("Monrovia crowds/streets 1985 (agency)", "https://www.gettyimages.com/photos/monrovia-1985", "ARCHIVE", "Getty - license"),
    "crowds": ("Crowd celebrating in a street (generic)", "https://www.pexels.com/search/crowd%20celebration/", "STOCK", "Pexels, free"),
    "radio": ("Vintage radio microphone / broadcast desk (generic)", "https://www.pexels.com/search/vintage%20microphone/", "STOCK", "Pexels, free"),
    "elbc": ("Radio station / transmitter mast (ELBC stand-in)", "https://www.pexels.com/search/radio%20station/", "STOCK", "Pexels, free"),
    "quiwonkpa": ("Thomas Quiwonkpa - SCARCE; only rare images exist", "https://en.wikipedia.org/wiki/Thomas_Quiwonkpa", "ARCHIVE", "Rare/likely uncleared - trace rights-holder. GAP"),
    "qd": ("Quiwonkpa & Doe / 1980 PRC junta together (archival)", "https://www.gettyimages.com/photos/samuel-doe", "ARCHIVE", "Getty - license"),
    "doe": ("Samuel Doe - portrait / public appearance", "https://commons.wikimedia.org/wiki/Category:Samuel_Doe", "CC", "Wikimedia - verify per-file; else Getty"),
    "reagan": ("Doe with Reagan at the White House, 1982", "https://catalog.archives.gov/search?q=Samuel%20Doe%20Reagan", "CC", "Reagan Library / US-gov likely PD - verify; else Getty"),
    "doedeath": ("Doe's 1990 capture/killing - GRAPHIC; agency/video", "https://www.gettyimages.com/photos/samuel-doe", "ARCHIVE", "Getty/AP - license; GRAPHIC, editorial caution"),
    "jackson": ("Jackson Doe (likely 1985 winner) - scarce", "https://en.wikipedia.org/wiki/Jackson_Doe", "ARCHIVE", "Rare - verify rights. GAP"),
    "taylor": ("Charles Taylor - portrait", "https://commons.wikimedia.org/wiki/Category:Charles_Taylor_(Liberian_politician)", "CC", "Wikimedia - verify per-file"),
    "johnson": ("Prince Johnson - portrait", "https://commons.wikimedia.org/wiki/Category:Prince_Johnson", "CC", "Wikimedia - verify; else Getty"),
    "sirleaf": ("Ellen Johnson Sirleaf (detained 1985) - free", "https://commons.wikimedia.org/wiki/Category:Ellen_Johnson_Sirleaf", "CC", "Wikimedia / US-gov - free"),
    "beach": ("The 1980 Monrovia BEACH EXECUTION of 13 ministers (famous photos)", "https://rarehistoricalphotos.com/cabinet-ministers-execution-liberia-1980/", "ARCHIVE", "AP/agency rights - license; GRAPHIC"),
    "coup80": ("1980 Liberian coup - soldiers / PRC (archival)", "https://www.gettyimages.com/photos/liberia-coup-1980", "ARCHIVE", "Getty - license"),
    "exec": ("Firing squad / military execution context (archival)", "https://www.gettyimages.com/photos/liberia-execution", "ARCHIVE", "Getty - license"),
    "qdeath": ("Quiwonkpa's body displayed 1985 - GRAPHIC, agency-held; consider a respectful alternative", "https://www.gettyimages.com/photos/quiwonkpa", "ARCHIVE", "Getty/AP - license; GRAPHIC - editorial caution"),
    "map": ("Liberia location map (Nimba, Monrovia marked)", "https://commons.wikimedia.org/wiki/File:Liberia_location_map.svg", "CC", "Wikimedia, free SVG"),
    "flag": ("Flag of Liberia", "https://commons.wikimedia.org/wiki/File:Flag_of_Liberia.svg", "CC", "Wikimedia, free"),
    "settler": ("19th-c. Americo-Liberian settlers / early Liberia (PD)", "https://commons.wikimedia.org/wiki/Category:History_of_Liberia", "CC", "Wikimedia - mostly PD by age; verify"),
    "truewhig": ("Americo-Liberian elite / True Whig era (archival/PD)", "https://en.wikipedia.org/wiki/Americo-Liberian_people", "CC", "Commons/LOC - PD by age; verify"),
    "ironore": ("Iron ore mining / LAMCO Nimba railway", "https://en.wikipedia.org/wiki/LAMCO", "CC", "Wikimedia - verify; else mining stock"),
    "firestone": ("Firestone rubber plantation, Harbel - rubber tapping", "https://www.propublica.org/article/firestone-and-the-warlord-chapter-2", "REAL", "ProPublica/agency - permission; or rubber-tapping stock"),
    "registry": ("Cargo ships / flag-of-convenience registry (generic)", "https://www.pexels.com/search/cargo%20ship/", "STOCK", "Pexels, free"),
    "voa": ("Voice of America transmitter / radio masts", "https://commons.wikimedia.org/wiki/Category:Voice_of_America", "CC", "Wikimedia - verify"),
    "coldwar": ("US flag / State Department (generic)", "https://www.pexels.com/search/american%20flag/", "STOCK", "Pexels, free"),
    "cash": ("US aid / stacks of dollars (generic)", "https://www.pexels.com/photo/paper-money-changing-hands-6694571/", "STOCK", "Pexels, free"),
    "ballots": ("Ballot box / vote counting (generic)", "https://www.pexels.com/search/ballot%20box/", "STOCK", "Pexels, free"),
    "capitol": ("US Capitol / Congress (free)", "https://commons.wikimedia.org/wiki/Category:United_States_Capitol", "CC", "Wikimedia / US-gov - free"),
    "border": ("Night border / jungle path (generic)", "https://www.pexels.com/search/jungle%20night/", "STOCK", "Pexels, free"),
    "soldiers": ("Soldiers / armed men silhouette (generic)", "https://pixabay.com/photos/silhouette-soldier-military-917961/", "STOCK", "Pixabay, free"),
    "monrovia": ("Monrovia cityscape (period or modern)", "https://www.gettyimages.com/photos/monrovia", "ARCHIVE", "Getty - license; some Commons free"),
    "night": ("City street at night / empty (generic)", "https://www.pexels.com/search/city%20street%20night/", "STOCK", "Pexels, free"),
    "mansion": ("Executive Mansion, Monrovia", "https://commons.wikimedia.org/wiki/Category:Executive_Mansion_(Liberia)", "CC", "Wikimedia - verify"),
    "somber": ("Somber transitional B-roll (overcast sky / empty street)", "https://www.pexels.com/search/overcast%20sky/", "STOCK", "Pexels, free"),
    "nimba": ("Nimba County landscape / village (context)", "https://www.gettyimages.com/photos/nimba-county", "ARCHIVE", "Getty - license; or HRW report"),
    "mine": ("Abandoned mine shaft (generic)", "https://www.pexels.com/search/mine%20shaft/", "STOCK", "Pexels, free"),
    "hrw": ("HRW 'Flight from Terror' report on Nimba abuses (1990)", "https://www.hrw.org/reports/1990/liberia2/", "DOC", "(c) HRW - permission for on-screen use"),
    "jail": ("Prison / detention (generic)", "https://pixabay.com/photos/prison-jail-barbed-wire-barbwire-482619/", "STOCK", "Pixabay, free"),
    "news": ("Shuttered newspaper / printing press (generic)", "https://www.pexels.com/search/newspaper%20printing/", "STOCK", "Pexels, free"),
    "refugees": ("Liberian refugees / displacement (archival)", "https://www.gettyimages.com/photos/liberia-refugees", "ARCHIVE", "Getty - license"),
    "war": ("First Liberian Civil War 1989-96 (archival)", "https://www.gettyimages.com/photos/liberian-civil-war", "ARCHIVE", "Getty/AP - license"),
    "exile": ("Lone figure / departure gate (exile, generic)", "https://www.pexels.com/search/airport%20silhouette/", "STOCK", "Pexels, free"),
    "trc": ("Liberia Truth & Reconciliation Commission report/hearings", "https://www.trcofliberia.org/", "REAL", "TRC - public; verify image rights"),
    "memorial": ("Modern Monrovia / reflective close", "https://www.gettyimages.com/photos/monrovia", "ARCHIVE", "Getty - license; some Commons free"),
    "title": ("Title card graphic to create", "—", "DOC", "Create"),
}

BEATS = [
    ("##", "COLD OPEN"),
    ("For one morning, it looked like the good guys had won.", "crowds85"),
    ("On the twelfth of November, 1985, the people of Monrovia woke to a voice on the radio", "radio"),
    ("they had not heard in two years.", "radio"),
    ("It belonged to Thomas Quiwonkpa -", "quiwonkpa"),
    ("the general who had once been the most popular man in Liberia,", "quiwonkpa"),
    ("before the dictator drove him into exile.", "doe"),
    ("He told them it was over. That he had come home.", "radio"),
    ("That Samuel Doe - the man who had stolen an election a month earlier - was finished.", "doe"),
    ("And the city believed him.", "crowds85"),
    ("People poured into the streets. They danced.", "crowds"),
    ("Some painted their bodies like warriors and waved palm branches.", "crowds"),
    ("After five years of fear, the dictatorship seemed to have fallen in a single morning.", "crowds85"),
    ("It hadn't.", "somber"),
    ("Within hours, the celebration would become a death sentence -", "somber"),
    ("for Quiwonkpa, for thousands of his people, and, in the end, for the country itself.", "nimba"),
    ("This is the story of the coup that almost worked.", "monrovia"),
    ("And of what it cost when it didn't.", "war"),
    ("[TITLE CARD] THE COUP THAT FAILED - LIBERIA, 1985", "title"),

    ("##", "ACT ONE - THE TWO SERGEANTS"),
    ("To understand that morning, you have to go back five years -", "map"),
    ("and Quiwonkpa and the man he came to overthrow had once been on the same side.", "qd"),
    ("Liberia was a country built on a fault line.", "map"),
    ("Founded in the 1820s by freed African-Americans,", "settler"),
    ("it was, for over a century, ruled by their descendants -", "truewhig"),
    ("a tiny settler elite called the Americo-Liberians -", "truewhig"),
    ("who governed the indigenous majority through a single party,", "truewhig"),
    ("and locked them out of the vote until 1951.", "settler"),
    ("For one hundred and thirty-three years, that was the order of things.", "settler"),
    ("Until April 1980.", "coup80"),
    ("That month, a group of low-ranking soldiers stormed the executive mansion,", "coup80"),
    ("killed the president, and ended settler rule at gunpoint.", "coup80"),
    ("Days later, they tied thirteen ministers to stakes on a Monrovia beach", "beach"),
    ("and shot them in front of the cameras.", "beach"),
    ("The man who emerged on top was a master sergeant named Samuel Doe -", "doe"),
    ("the first indigenous Liberian ever to rule the country.", "doe"),
    ("And at his side, a hero of the revolution, was a young general named Thomas Quiwonkpa.", "quiwonkpa"),
    ("Quiwonkpa was everything Doe was not.", "quiwonkpa"),
    ("From Nimba County, an ethnic Gio, he was charismatic,", "quiwonkpa"),
    ("widely trusted, seen by many as the conscience of the new government.", "quiwonkpa"),
    ("And that, it turned out, was a dangerous thing to be.", "doe"),
    ("Because Samuel Doe did not intend to share power.", "doe"),
    ("One by one, he turned on the comrades who had made him.", "coup80"),
    ("In 1981, his deputy was executed by firing squad.", "exec"),
    ("By 1983, Quiwonkpa - too popular, too principled -", "quiwonkpa"),
    ("was pushed out, accused of plotting, and forced to flee the country.", "exile"),
    ("Doe increasingly ruled through his own small ethnic group, the Krahn -", "doe"),
    ("packing the army with them,", "soldiers"),
    ("and treating Quiwonkpa's people, the Gio and the Mano, as enemies in waiting.", "nimba"),
    ("The revolution had eaten itself.", "coup80"),
    ("And its most popular son was in exile, watching.", "exile"),

    ("##", "ACT TWO - WHY DOE COULD NOT FALL"),
    ("There was a reason Samuel Doe felt untouchable. Two reasons, really.", "doe"),
    ("The first was money. The second was Washington.", "cash"),
    ("Liberia was not rich, but it had things the world wanted.", "ironore"),
    ("Iron ore - for a time, more than half the country's exports -", "ironore"),
    ("dug out of the Nimba mountains.", "ironore"),
    ("And rubber: the vast Firestone plantation at Harbel, the largest in the world,", "firestone"),
    ("an American corporation that had operated here for half a century.", "firestone"),
    ("Whoever held the capital held the concession money, the shipping registry,", "registry"),
    ("and the patronage that kept an army loyal.", "soldiers"),
    ("But the deeper reason was the Cold War.", "coldwar"),
    ("To the United States, tiny Liberia was a strategic prize:", "voa"),
    ("home to powerful Voice of America transmitters, a navigation station,", "voa"),
    ("and the main CIA listening post in Africa.", "coldwar"),
    ("Washington wanted a reliable anti-communist friend in West Africa -", "coldwar"),
    ("and Samuel Doe was willing to be one.", "reagan"),
    ("So the money flowed.", "cash"),
    ("Liberia became the largest per-person recipient of American aid in sub-Saharan Africa.", "cash"),
    ("Doe was invited to the White House. Twice.", "reagan"),
    ("Which is why, in October 1985, when Doe held an election and lost it -", "ballots"),
    ("and then counted the ballots in secret and declared himself the winner with fifty-point-nine percent -", "ballots"),
    ("almost nothing happened.", "doe"),
    ("The man most observers believed had actually won, Jackson Doe, no relation,", "jackson"),
    ("was brushed aside.", "jackson"),
    ("The United States Congress called the election fraudulent.", "capitol"),
    ("The Reagan administration called it a step toward democracy.", "reagan"),
    ("And the aid kept coming.", "cash"),
    ("For Liberians who had hoped to vote their way out, the message was clear.", "crowds"),
    ("There was no peaceful way out.", "somber"),
    ("One month later, Thomas Quiwonkpa came home.", "quiwonkpa"),

    ("##", "ACT THREE - THE MORNING IT ALMOST WORKED"),
    ("He came quietly, across the border from Sierra Leone,", "border"),
    ("with a small band of fighters - by most accounts, only about two dozen men.", "soldiers"),
    ("It was almost nothing. And for a few hours, it was almost enough.", "monrovia"),
    ("In the early hours of the twelfth of November,", "radio"),
    ("his men seized the national radio station and the main army barracks.", "elbc"),
    ("They arrested several of Doe's ministers.", "soldiers"),
    ("And Quiwonkpa went on the air and told the country he had taken power -", "radio"),
    ("that Doe was in hiding, and there was no escape for him.", "doe"),
    ("The city erupted.", "crowds85"),
    ("For those few hours, Quiwonkpa's forces controlled Monrovia,", "monrovia"),
    ("and the people celebrated as if the war was already won.", "crowds"),
    ("But it wasn't won.", "somber"),
    ("And here the story turns on a single, fatal mistake.", "somber"),
    ("Quiwonkpa announced his victory before he had confirmed it.", "radio"),
    ("He told the nation Doe was finished -", "radio"),
    ("but Doe was not dead, and Doe was not captured.", "doe"),
    ("He was simply hiding. And a hidden enemy is not a defeated one.", "doe"),
    ("That gap - the difference between almost and done -", "somber"),
    ("was all Samuel Doe needed.", "doe"),
    ("By mid-afternoon, he surfaced on the same radio waves to say he was still in charge.", "doe"),
    ("His Krahn soldiers, the Executive Mansion Guard, regrouped and counterattacked.", "soldiers"),
    ("And Quiwonkpa's tiny force began to come apart.", "soldiers"),
    ("He lost contact with his own men.", "somber"),
    ("Loyalist troops broke into the stockade and freed the ministers.", "mansion"),
    ("The crowds that had danced in the morning melted away.", "somber"),
    ("Two dozen men, however brave, could not hold a capital against a regrouping army.", "soldiers"),
    ("By nightfall, the coup that had owned Monrovia at dawn was finished.", "night"),
    ("Three days later, on the fifteenth of November,", "somber"),
    ("soldiers cornered Thomas Quiwonkpa in a house outside the city, and killed him.", "quiwonkpa"),
    ("What they did next, they did deliberately, and in public.", "somber"),
    ("His body was mutilated and put on display.", "qdeath"),
    ("And by many accounts - consistently reported, though never proven -", "qdeath"),
    ("soldiers ate parts of it, believing his power would pass into them.", "qdeath"),
    ("It was meant as a message. The message was received.", "somber"),

    ("##", "ACT FOUR - THE RECKONING"),
    ("If the story ended with Quiwonkpa's death, it would only be a tragedy.", "somber"),
    ("What makes it a catastrophe is what came after.", "nimba"),
    ("Samuel Doe did not stop at the men who had attacked him.", "doe"),
    ("He turned on Quiwonkpa's people - the Gio and the Mano of Nimba County -", "nimba"),
    ("most of whom had nothing to do with the coup at all.", "nimba"),
    ("His forces, led by the head of his mansion guard, swept through Nimba:", "soldiers"),
    ("rounding up civilians, executing them, throwing bodies down a disused mine shaft.", "mine"),
    ("How many died, no one can say precisely -", "nimba"),
    ("the estimates run from hundreds into the thousands.", "nimba"),
    ("Human rights investigators documented it as systematic, ethnic killing.", "hrw"),
    ("In Monrovia, Doe arrested the opposition leaders -", "jail"),
    ("including the man who had probably won the election -", "jackson"),
    ("jailed journalists, and shut the newspapers.", "news"),
    ("And in doing all of it, he lit a fuse.", "war"),
    ("Because the Gio and Mano survivors of Nimba did not forget.", "refugees"),
    ("Four years later, a warlord named Charles Taylor crossed into that same county,", "taylor"),
    ("and thousands of them joined him - not for ideology, but for revenge.", "war"),
    ("That rebellion became the First Liberian Civil War.", "war"),
    ("Before it was over, more than a quarter of a million people would be dead.", "war"),
    ("Samuel Doe did not survive it.", "doe"),
    ("In 1990, he was captured by a former comrade of Quiwonkpa's -", "johnson"),
    ("and tortured to death, the killing filmed and broadcast,", "doedeath"),
    ("his body mutilated, just as Quiwonkpa's had been five years before.", "doedeath"),
    ("The circle had closed.", "somber"),

    ("##", "CLOSE"),
    ("Years later, Liberia's own Truth and Reconciliation Commission", "trc"),
    ("traced the long war back through the killings in Nimba,", "nimba"),
    ("to that single failed morning in November 1985.", "radio"),
    ("It is tempting to remember the coup that failed as a near miss -", "monrovia"),
    ("the moment Liberia almost turned a corner.", "crowds85"),
    ("But the harder truth is the one the country lived:", "war"),
    ("that the morning it almost worked, and the way it didn't, set in motion fourteen years of war.", "war"),
    ("Thomas Quiwonkpa gambled that a popular man with a couple dozen soldiers could topple a dictator.", "quiwonkpa"),
    ("For one morning, he was almost right.", "crowds85"),
    ("And the price of almost was paid, in the end, by everyone.", "memorial"),
]


def fmt(t):
    return f"{int(t // 60):02d}:{int(t % 60):02d}"


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)
for lvl, sz in [(1, 19), (2, 14)]:
    doc.styles[f"Heading {lvl}"].font.color.rgb = NAVY
    doc.styles[f"Heading {lvl}"].font.size = Pt(sz)

doc.add_heading("THE COUP THAT FAILED (LIBERIA, 1985) - Timed Voiceover & Shot List", level=1)
doc.add_paragraph().add_run("Source: liberia_1985_script.md (Node 2)  |  one visual per ~3-5s beat.").italic = True
legend = doc.add_paragraph()
legend.add_run("Asset legend:  ").bold = True
for k, label in [("REAL", "real/specific"), ("CC", "public-domain/CC (free)"), ("ARCHIVE", "licensable archive"), ("STOCK", "generic stock"), ("DOC", "document/graphic")]:
    r = legend.add_run(f"{k} "); r.bold = True; r.font.color.rgb = TYPE_COLOR[k]
    legend.add_run(f"= {label}   ").font.size = Pt(9)
d = doc.add_paragraph()
d.add_run("Rights note: the footage sub-agents hit a session limit, so this library was assembled "
          "directly from reputable sources - links are a starting map to VERIFY, not cleared rights. "
          "Most 1980s Liberia event footage is AP/Reuters/Getty and must be licensed. Several beats are "
          "GRAPHIC (Quiwonkpa's body; Doe's 1990 death) - handle with editorial care. Priority: real -> "
          "public-domain/CC -> licensable archive -> stock.").italic = True
doc.add_paragraph()

clock = 0.0
md = ["# THE COUP THAT FAILED (LIBERIA, 1985) - Timed Voiceover & Shot List\n",
      "_One visual per ~3-5s beat. Links = starting map to verify, not cleared rights._\n"]
for text, key in BEATS:
    if text == "##":
        doc.add_heading(key, level=2); md.append(f"\n## {key}\n"); continue
    words = len(text.split())
    start, end = clock, clock + words / WPM * 60
    clock = end
    tcode = f"[{fmt(start)}-{fmt(end)}]"
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
    rt = p.add_run(tcode + "  "); rt.bold = True; rt.font.color.rgb = TC; rt.font.size = Pt(9)
    p.add_run(text)
    desc, link, typ, lic = A[key]
    v = doc.add_paragraph(); v.paragraph_format.left_indent = Inches(0.3); v.paragraph_format.space_after = Pt(6)
    vr = v.add_run("VISUAL: "); vr.bold = True; vr.font.size = Pt(9); vr.font.color.rgb = NAVY
    v.add_run(desc + "  ").font.size = Pt(9)
    tr = v.add_run(f"[{typ}] "); tr.bold = True; tr.font.size = Pt(9); tr.font.color.rgb = TYPE_COLOR[typ]
    lr = v.add_run(f"{link}  ({lic})"); lr.font.size = Pt(8); lr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    md.append(f"{tcode} {text}")
    md.append(f"   - VISUAL: {desc} [{typ}] {link} ({lic})\n")

n = sum(1 for t, _ in BEATS if t != "##")
doc.add_paragraph()
doc.add_paragraph(f"Total beats: {n}  |  Computed runtime: {fmt(clock)}").italic = True
doc.save("liberia_1985_shotlist.docx")
open("liberia_1985_shotlist.md", "w").write("\n".join(md) + f"\n\n_Total beats: {n} | runtime {fmt(clock)}_\n")
print(f"Beats: {n}  runtime {fmt(clock)}  -> liberia_1985_shotlist.docx / .md")
