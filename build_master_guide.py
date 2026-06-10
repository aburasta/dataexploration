#!/usr/bin/env python3
"""Build the consolidated PROJECT MASTER GUIDE (.docx) for upload to Claude Projects.
Folds together: concept, assets, visual style, per-episode workflow, prompt pack,
Higgsfield operations (models/credits/consistency), procedure, and checklists."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FULL_SUFFIX = (
    "Bright, cheerful, wholesome flat 2D vector children's-book clip-art illustration. Rounded, cute, friendly "
    "characters and objects with clean simple shapes and chunky proportions. Soft subtle shading and gentle soft "
    "highlights for a little depth; thin minimal darker edge accents only where a shape needs separation (no heavy "
    "black outlines). Small dot eyes, rosy cheek blush, simple happy expressions. Bright multicolor brand palette: "
    "sunshine yellow (#FBC02D), sky blue (#4FC3E8), leaf green (#7CC576), tomato coral (#F0654E), tangerine "
    "(#F39237), blush pink (#F6A9B6), stone gray (#9DB0BC), warm brown (#8D6E63), cream (#FBF7EF), charcoal "
    "(#3A3631). SINGLE isolated element only, centered, on a plain TRANSPARENT background — no scene, no setting, no "
    "ground shadow. Clean, modern, wholesome. No grain, no gradients, no 3D, no realistic textures, no photorealism, "
    "no heavy black outlines, no clutter, no scene background."
)
COMPACT_SUFFIX = (
    "STYLE: bright cheerful flat 2D vector children's-book clip-art; rounded cute friendly shapes, soft subtle "
    "shading + gentle highlights, thin minimal edge accents (no heavy black outlines); dot eyes, rosy cheeks; "
    "palette sunshine #FBC02D, sky #4FC3E8, leaf #7CC576, coral #F0654E, tangerine #F39237, blush #F6A9B6, stone "
    "#9DB0BC, brown #8D6E63, cream #FBF7EF, charcoal #3A3631; SINGLE isolated element, centered, TRANSPARENT "
    "background, no scene/no ground shadow; clean modern wholesome. No grain/gradient/3D/realistic/photoreal/heavy "
    "outline/clutter/scene-background."
)


def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexv.lstrip("#"))
    tcPr.append(shd)


def sct(cell, text, bold=False, size=9, color=None):
    cell.text = ""; r = cell.paragraphs[0].add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def h1(t): return doc.add_heading(t, level=1)
def h2(t): return doc.add_heading(t, level=2)


def body(t, size=10.5, italic=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.size = Pt(size); r.italic = italic
    return p


def bullet(t, size=10.5):
    p = doc.add_paragraph(style="List Bullet"); p.add_run(t).font.size = Pt(size); return p


def num(t, size=10.5):
    p = doc.add_paragraph(style="List Number"); p.add_run(t).font.size = Pt(size); return p


def box(label, text):
    if label:
        lp = doc.add_paragraph(); r = lp.add_run(label); r.bold = True
        r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string("1F4E5F")
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); shade(c, "F2F4F0"); c.text = ""
    rr = c.paragraphs[0].add_run(text); rr.font.name = "Consolas"; rr.font.size = Pt(9)
    doc.add_paragraph()


def callout(text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    c = tbl.cell(0, 0); shade(c, "FBEFD6"); c.text = ""
    r = c.paragraphs[0].add_run(text); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("8A4B16"); doc.add_paragraph()


doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

# ---------------- TITLE ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PROJECT MASTER GUIDE"); r.bold = True; r.font.size = Pt(28)
r.font.color.rgb = RGBColor.from_string("1F4E5F")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Finance × Animal Behavior — Animated YouTube Documentary Series")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string("EF5A2A")
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("The single, complete reference for the project. Upload to the Claude Projects section. "
              "When a spreadsheet row is named, follow this guide end-to-end.")
r.italic = True; r.font.size = Pt(10)
doc.add_paragraph()

# ---------------- 1. AT A GLANCE ----------------
h1("1. The Project at a Glance")
body("A YouTube documentary series that explains difficult finance concepts by paralleling them to specific animal "
     "behaviors. Each episode is built from ONE row of the file finance_animal_behavior_matches.xlsx.")
body("The pipeline for every episode:")
num("Pick a row (a finance concept + its animal-behavior match).")
num("Write the narration script (benefit-led hook + retention arc).")
num("Break it into a timed SHOT LIST — a new on-screen visual every 3–5 seconds.")
num("Produce a HIGGSFIELD PROMPT PACK — paste-ready prompts that generate every ELEMENT in the fixed style.")
num("Generate each frame's elements as separate TRANSPARENT PNGs (Higgsfield + background removal), named per frame.")
num("Assemble in VideoScribe: the animated hand DRAWS each element on, frame by frame; add VO + captions, then publish.")
callout("TWO THINGS NEVER CHANGE BETWEEN EPISODES: the VISUAL STYLE (Section 4) and the SCRIPT/HOOK FORMULA "
        "(Section 3). Only the content — the concept and the animal — changes. Consistency is the brand.")
bullet("Target runtime: 10–15 minutes (default ~12).  Cadence: a visual change every 3–5 seconds.")
bullet("A 10–15 min episode = ~150–300 frames; each frame = several isolated element PNGs drawn on in sequence.")
bullet("Primary image model: Seedream 4.5.  Assembly + motion: VideoScribe animated hand-reveal.")

# ---------------- 2. ASSETS ----------------
h1("2. Project Assets & Inputs")
tbl = doc.add_table(rows=1, cols=2); tbl.style = "Table Grid"
for i, x in enumerate(["File", "What it is / how it's used"]):
    sct(tbl.rows[0].cells[i], x, bold=True, size=9, color="FFFFFF"); shade(tbl.rows[0].cells[i], "1F4E5F")
for f, d in [
    ("finance_animal_behavior_matches.xlsx", "THE INPUT. 100 granular finance concepts, each with 2 rated animal "
     "matches + script-ready behavior descriptions. One row = one episode. Default to the higher-rated match."),
    ("finance_concepts_for_animation.xlsx", "The source list of 100 granular finance concepts (reference)."),
    ("PROJECT_MASTER_GUIDE.docx", "THIS document — the complete project guide."),
    ("style reference image", "The approved cartoon look (bright children's-book clip-art) — see Section 4."),
    ("VideoScribe (Sparkol)", "THE ASSEMBLY TOOL. The animated hand draws each element PNG onto a 16:9 canvas, "
     "frame by frame, to build the video (see the VideoScribe Assembly section)."),
    ("EP07_risk-pooling_vampire-bat/", "A complete worked demo episode (shot list + prompt pack) to copy as a template."),
]:
    c = tbl.add_row().cells
    sct(c[0], f, bold=True, size=9); sct(c[1], d, size=9)
tbl.rows[0].cells[0].width = Inches(2.6); tbl.rows[0].cells[1].width = Inches(4.0)
doc.add_paragraph()
body("Reading a row: use the Category, Granular Concept, and The Specific Finance Aspect to know exactly what to "
     "teach; use the chosen Animal Match's description as the factual basis. If a match is flagged below 75% or "
     "'debated', lean into the 'where the analogy breaks down' beat rather than overstating it.")

# ---------------- 3. CREATIVE PRINCIPLES ----------------
h1("3. Creative Principles — Script & Hook (fixed every episode)")
h2("3a. The Hook (first ~10–15 seconds)")
body("ALWAYS create suspense around the personal BENEFIT of knowing the lesson — framed as an intriguing question or "
     "'what if'. Do NOT open by naming the concept; open with the payoff and the mystery.")
box("HOOK EXAMPLES (benefit-led):",
    "Avoid herd behavior → \"How do the people who spot the next big thing always see it first? "
    "The answer is hiding in how one bird breaks from the flock...\"\n"
    "Risk pooling → \"What if one bad month could never wipe you out — no matter how unlucky you got?\"\n"
    "Compounding → \"What if one small, boring habit today decided whether you're rich in 20 years?\"")
h2("3b. Retention Arc (after the hook)")
for x in [
    "STAKES — make the viewer feel why getting this wrong hurts them.",
    "MEET THE ANIMAL — introduce the animal + behavior vividly (from the spreadsheet description).",
    "THE PARALLEL — map the behavior onto the finance aspect, one idea at a time.",
    "WHERE IT BREAKS DOWN — honestly note where the analogy stops (signature trust-builder).",
    "THE PAYOFF — deliver the practical lesson promised in the hook.",
    "BUTTON — a short memorable closing line + soft subscribe / next-episode nudge.",
]:
    bullet(x)
bullet("Throughout: short sentences, concrete imagery, a question every ~20–30 seconds, warm and clever — never lecture-y.")
h2("3c. Length & Pace")
bullet("Target 10–15 minutes (default ~12). Narration pace 150 wpm = 2.5 words/sec.")
bullet("At 150 wpm: 10 min ≈ 1,500 words, 12 min ≈ 1,800, 15 min ≈ 2,250.")

# ---------------- 4. VISUAL STYLE & COLOR CONSISTENCY ----------------
h1("4. Visual Style & Color Consistency — Applied to EVERY Episode")
callout("THE ONE RULE: the SAME palette + the SAME style on every frame of every episode — that sameness IS the "
        "brand. Only the CONTENT changes between videos. Append a style tag (4e / Appendix) to every prompt; if a "
        "new image doesn't match this section, regenerate it. Every element is generated ISOLATED on a transparent "
        "background.")
h2("4a. The style in one sentence")
body("Bright, cheerful, wholesome FLAT 2D vector children's-book CLIP-ART: rounded, cute, friendly characters and "
     "objects; clean simple shapes; soft subtle shading with gentle highlights; thin minimal edge accents (no heavy "
     "black outlines); small dot eyes and rosy cheeks. Every asset is a SINGLE isolated element on a transparent "
     "background, drawn onto the cream VideoScribe canvas.")
h2("4b. Style rules (the non-negotiables)")
for x in [
    "CLIP-ART, ISOLATED — one element per asset, centered, on a TRANSPARENT background. No scenes, no settings, no ground shadow.",
    "ROUNDED & CUTE — chunky friendly proportions; circles, ovals, soft rounded shapes.",
    "SOFT SHADING + HIGHLIGHTS — gentle subtle depth; never flat-dead, never 3D.",
    "THIN EDGE ACCENTS ONLY — light darker edges where shapes need separation; NO heavy black ink outlines.",
    "BRAND PALETTE ONLY (4c) — bright and cheerful, but fixed. Re-tint real animals into the palette.",
    "FRIENDLY FACES — small dot eyes, rosy cheek blush, simple happy expressions.",
    "CLEAN & WHOLESOME — uncluttered, modern; NO grain, NO texture overlays, NO gradients.",
    "ONE FOCAL COLOR — reserve TOMATO CORAL for the single thing the viewer should look at in each frame.",
]:
    bullet(x)
h2("4c. The color palette (canonical — use these exact hex values)")
palette = [
    ("Sunshine Yellow", "#FBC02D", "Beaks, manes, warm highlights"),
    ("Sky Blue", "#4FC3E8", "Water/sky elements, calm characters, the Host"),
    ("Leaf Green", "#7CC576", "Foliage, ground shapes"),
    ("Tomato Coral", "#F0654E", "THE focal color — the ONE key idea per frame (sparingly)"),
    ("Tangerine", "#F39237", "Secondary warm accent (lion/monkey)"),
    ("Blush Pink", "#F6A9B6", "Cheek blush, soft accents"),
    ("Stone Gray", "#9DB0BC", "Elephant/rhino bodies, neutral props"),
    ("Warm Brown", "#8D6E63", "Trunks, wood, branches"),
    ("Cream", "#FBF7EF", "The VideoScribe canvas / negative space"),
    ("Charcoal", "#3A3631", "Eyes, edge accents, fine detail — never pure black"),
]
pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
for i, x in enumerate(["Swatch", "Name", "Hex", "Where it's used"]):
    sct(pt.rows[0].cells[i], x, bold=True, size=8.5, color="FFFFFF"); shade(pt.rows[0].cells[i], "1F4E5F")
for name, hexv, usage in palette:
    c = pt.add_row().cells
    shade(c[0], hexv); sct(c[0], "", size=8.5)
    sct(c[1], name, bold=True, size=8.5); sct(c[2], hexv, size=8.5); sct(c[3], usage, size=8.5)
for row in pt.rows:
    row.cells[0].width = Inches(0.5); row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(0.9); row.cells[3].width = Inches(3.3)
doc.add_paragraph()
bullet("Bright but balanced: spread the palette across elements; let TOMATO CORAL spotlight the one key idea.")
bullet("Cream (#FBF7EF) is the canvas/background. No pure black/white — use Charcoal and Cream. Tints/shades of "
       "these colors are fine; brand-new colors are not.")
h2("4d. Character anchors (describe with the SAME words every time)")
body("Re-stating these exact descriptions in every prompt is what keeps recurring characters identical across all "
     "videos. Save the Host (and any repeated prop) as a Higgsfield ELEMENT and reuse it — don't redraw it.")
box("HOST (recurring narrator — in every episode):",
    "cute clip-art fox host (rounded chunky sky-blue body, charcoal dot eyes, rosy cheeks, tomato-coral knitted "
    "scarf). Lock ONE design and reuse it identically as a saved Element.")
box("ANIMAL OF THE WEEK (the episode's animal):",
    "the episode's animal as a cute clip-art character, RE-TINTED into the brand palette (NOT natural photo colors), "
    "rounded and friendly, simplified, dot eyes, rosy cheeks.")
box("CROWDS / MARKETS / 'YOU':",
    "simple rounded clip-art person figure (sky-blue body, charcoal dot eyes, no other facial detail); a crowd is "
    "many identical figures, the odd-one-out is tomato coral.")
bullet("FINANCE MOTIFS: charts/coins/arrows as cute chunky clip-art in the palette (no stock screenshots); keep one "
       "icon vocabulary; each is its OWN isolated transparent element.")
bullet("COMPOSITION: element PNGs have no fixed aspect ratio; the 16:9 LAYOUT is built on the cream VideoScribe "
       "canvas — one clear focal subject per frame, bottom ~15% caption-safe.")
h2("4e. Style tags (paste at the end of EVERY prompt)")
body("Append one verbatim so the look is locked: the COMPACT tag for normal per-element prompts; the FULL suffix "
     "for one-off hero elements. Full text of both, plus the NEGATIVE/AVOID list, is in the Appendix.")
box("COMPACT STYLE TAG (the workhorse — append to every per-element prompt):", COMPACT_SUFFIX)

# ---------------- 5. DELIVERABLES + TIMING ----------------
h1("5. Per-Episode Deliverables")
body("For each row, produce THREE linked deliverables (Frame 001 ↔ its element prompts ↔ frame_001/ element PNGs):")
num("DELIVERABLE A — SHOT LIST (Word doc): header meta + clean narration script + a table of timed FRAMES "
    "(Frame #, timestamp, narration slice, one-line visual, ELEMENT BREAKDOWN, draw order, prompt #s, STILL/CLIP).")
num("DELIVERABLE B — HIGGSFIELD PROMPT PACK (text): a consistency header + per-ELEMENT prompts grouped by frame "
    "(Section 6) — one paste-ready prompt for each isolated element.")
num("DELIVERABLE C — FRAMES FOLDER: for each frame, every element as a TRANSPARENT PNG, produced by running the "
    "Prompt Pack + background removal, named frame_0NN_elXX_<name>.png inside frames/frame_0NN/.")
h2("5a. Timing method (every frame = 3–5 seconds)")
num("Write the full clean narration script.")
num("Total duration = total_words ÷ 2.5 (seconds).")
num("Segment into frames of 3–5s each (~8–13 words = one short clause), one on-screen scene per frame.")
num("List each frame's ELEMENTS and the order the hand draws them; assign cumulative timestamps (frame length = "
    "its_words ÷ 2.5). The frame's elements are drawn on in sequence across its narration slice.")
callout("RESULT: a 10–15 min episode → ~150–300 frames, each composed of a few isolated element PNGs. Reuse saved "
        "ELEMENTS (Host, recurring props) across frames so you don't regenerate the same asset every time.")
h2("5b. File/folder naming")
box("FOLDER STRUCTURE:",
    "EP07_risk-pooling_vampire-bat/\n"
    "    EP07_shot-list.docx            <- Deliverable A\n"
    "    EP07_prompt-pack.txt           <- Deliverable B\n"
    "    frames/                        <- Deliverable C (transparent element PNGs)\n"
    "        frame_001/  frame_001_el01_<name>.png  frame_001_el02_<name>.png ...\n"
    "        frame_002/  ...\n"
    "    videoscribe/  EP07.scribe      <- the VideoScribe assembly project\n"
    "    clips/   frame_0XX.mp4         <- OPTIONAL image-to-video for special motion")

# ---------------- 6. PROMPT PACK ----------------
h1("6. The Higgsfield Prompt Pack (Deliverable B)")
body("After the shot list, ALWAYS output the Prompt Pack: the exact prompts the user pastes into the Higgsfield web "
     "app to generate the episode's ELEMENTS consistently. Each frame becomes a small group of element prompts. "
     "Two parts:")
h2("6a. The Consistency Header (written once)")
bullet("STYLE BLOCK: the compact style tag (below), verbatim.")
bullet("CHARACTER BIBLE: thorough, word-for-word descriptions of every recurring character (Host, Animal of the "
       "Week, crowd figures). Reuse these EXACT words in every prompt featuring the character.")
bullet("PROP/MOTIF BIBLE: the recurring props and finance motifs (coins, charts, arrows), described identically.")
bullet("PALETTE hex list + SETTINGS (Seedream 4.5, attach style key-frame as reference + Host Element, one seed per "
       "character/prop, ISOLATED element on a plain background → remove background to a transparent PNG).")
h2("6b. The per-element prompts (grouped by frame, numbered to the shot list)")
bullet("Group prompts under each Frame # and list the DRAW ORDER (the order the hand reveals them on the canvas).")
bullet("Each = Consistency Header content + [the ONE isolated element] + 'single isolated element, centered, "
       "transparent background, no scene, no ground shadow' + the COMPACT style tag.")
bullet("ONE element per prompt — a character, a prop, or a motif. RE-STATE recurring characters using the SAME "
       "wording as the Character Bible — never 'same as before' (that causes drift).")
bullet("Reuse a saved ELEMENT instead of re-prompting the Host/recurring props where possible; state the seed.")
bullet("STAY UNDER 3000 CHARACTERS (Higgsfield's per-prompt limit). Use the COMPACT style tag, not the full "
       "paragraph — each self-contained prompt should land ~500–900 characters.")
callout("THOROUGHNESS RULE: every per-element prompt carries a COMPLETE description of the element, palette, and "
        "style, and asks for a single isolated element on a transparent background — that is what keeps every "
        "separately-generated asset on-model and clean to layer in VideoScribe.")
box("COMPACT STYLE TAG (append to every per-element prompt):", COMPACT_SUFFIX)
box("EXAMPLE PER-ELEMENT PROMPT (~600 chars):",
    "Single isolated clip-art vampire bat — rounded chunky body in stone gray with sky-blue wing membranes, tiny "
    "charcoal dot eyes, small white fangs, rosy cheeks, friendly, hanging upside down, centered. Single isolated "
    "element, transparent background, no scene, no ground shadow. " + COMPACT_SUFFIX)

# ---------------- 7. HIGGSFIELD EXECUTION (TRANSPARENT ELEMENT PNGs) ----------------
h1("7. Higgsfield Execution — Generating Transparent Element PNGs")
body("This is the exact loop for turning each element prompt (Section 6) into a clean, transparent PNG ready to drop "
     "onto the VideoScribe canvas. Do this for EVERY element of EVERY frame (reusing saved Elements where you can).")
h2("7a. Generate the isolated element")
num("Pick the model: Seedream 4.5 (seedream_v4_5) for the bulk — UNLIMITED/free in the web app on a paid plan. Use "
    "Nano Banana Pro only for elements that contain text/numbers (it renders text better).")
num("Paste the element prompt: ONE isolated subject, centered, described as 'single isolated element on a plain "
    "background' + the COMPACT style tag. Attach the STYLE KEY-FRAME as a reference and the HOST/prop ELEMENT when "
    "the element is a recurring character or prop.")
num("Generate 2–4 variations and keep the cleanest, most on-model one. Keep one SEED per character/prop so the same "
    "asset looks identical across frames.")
h2("7b. Cut it out to a transparent PNG")
bullet("Seedream/Nano Banana FILL the background — so the cutout step is what gives you the alpha. Generate on a "
       "plain, flat, high-contrast background (e.g. plain cream or plain blue) to make removal clean.")
bullet("Run BACKGROUND REMOVAL on the chosen image: in the web app use the remove-background tool, or via MCP call "
       "remove_background with the generation's media_id and media_type='image'. Result = a transparent cutout.")
bullet("Download the cutout and save it as frames/frame_0NN/frame_0NN_elXX_<name>.png (PNG with alpha).")
callout("WHY TRANSPARENT: isolated alpha PNGs layer cleanly on the VideoScribe canvas with no white halos, so the "
        "hand can draw each element on independently and they overlap naturally to build the frame.")
h2("7c. Optional — SVG line version (true pen strokes)")
bullet("Flat PNGs REVEAL under the moving hand (a wipe-on). For genuine pen-stroke drawing, also make an SVG line "
       "version of the element (auto-trace the PNG in Illustrator Image Trace / Inkscape) and let the flat color "
       "fill behind it in VideoScribe. Use for hero elements where the drawing motion should really read.")
callout("Section 8 covers models/credits/clips; the VideoScribe Assembly section covers putting these PNGs together.")

# ---------------- 8. HIGGSFIELD OPERATIONS ----------------
h1("8. Higgsfield Operations (models, consistency, credits)")
h2("8a. Models")
bullet("PRIMARY: Seedream 4.5 (id seedream_v4_5) — high quality, up to 4K, ~1 credit/image, in the paid plans' "
       "365-day unlimited set. Default for ALL elements. quality 'high' for hero elements, 'basic' otherwise.")
bullet("SECONDARY: Nano Banana Pro or GPT Image — ONLY for elements with on-image text/diagrams (better text). "
       "Cost more; use sparingly. (Better still: add text in VideoScribe / the editor.)")
h2("8b. Web app vs API (this decides whether generation is FREE)")
bullet("Higgsfield WEB APP on a paid plan: Seedream is UNLIMITED (free) — generate the bulk here.")
bullet("Via API/automation (e.g., generating inside a chat): Seedream still bills ~1 credit/image even on a paid plan.")
bullet("FREE tier: ~10 credits/day, reloading daily; cannot sustain a full episode (use it for style tests only).")
h2("8c. Consistency tools (so every element stays on-model)")
bullet("Save the HOST + recurring props as ELEMENTS (for Seedream use Elements + reference images — NOT Soul) and "
       "REUSE them across frames instead of regenerating.")
bullet("Keep one approved STYLE KEY-FRAME and attach it as a reference image to every generation.")
bullet("Reuse ONE SEED per character/prop so the same asset is identical in every frame it appears.")
bullet("Generate 2–4 variations per element and keep the best — AI drifts slightly even with anchors.")
h2("8d. Clips (optional motion)")
bullet("Default motion is the VideoScribe hand-reveal. Use a short (2–4s) image-to-video CLIP only where real motion "
       "is essential (a coin flip, a number ticking, a chart drawing). Generate the element still first, then "
       "animate it; drop the clip onto the VideoScribe/editor timeline at that frame.")
h2("8e. Plans & credits")
bullet("Seedream ≈ 1 credit/image; Nano Banana Pro ≈ 2.  Free reloads ~10/day (no rollover, basic models only).")
bullet("Plus ≈ $39/mo annual = 1,000 credits/mo (~2–4 episodes on credits alone).")
bullet("Ultra ≈ $99/mo annual = 3,000 credits/mo + best unlimited coverage + lowest cost/credit (recommended for "
       "daily output, generating the bulk free in the web app).")
bullet("Credits expire 90 days and don't roll over. Daily output on credits alone would need ~9,000–18,000/mo — so "
       "rely on web-app unlimited Seedream, not credits.")

# ---------------- 9. VIDEOSCRIBE ASSEMBLY ----------------
h1("9. VideoScribe Assembly (Draw-On Reveal)")
body("This is where the episode is built: the animated hand draws each element PNG onto the canvas, frame by frame, "
     "to produce the video. Recommended tool: VideoScribe (Sparkol); the desktop app handles imported PNGs best.")
h2("9a. Project setup (once per episode)")
num("New scribe → set the canvas to 16:9 and the background to CREAM (#FBF7EF) — the brand canvas.")
num("Decide voiceover first (import VO audio) so element draw-times can be matched to the narration.")
h2("9b. Build each frame (in narration order)")
num("Import the frame's element PNGs (the image/camera icon → Add image) from frames/frame_0NN/.")
num("Position and scale each element to its place in the 16:9 layout; keep the bottom ~15% caption-safe.")
num("Per element, open its properties and set DRAW TIME (e.g. 1–3s), PAUSE, and TRANSITION so the elements reveal "
    "in the planned DRAW ORDER across the frame's narration slice (frame total ≈ 3–5s).")
num("Choose a drawing HAND (or 'no hand' for a clean wipe-on). Use the same hand throughout for consistency.")
num("On the timeline, order all elements of the frame, then move to the next frame; repeat. The sequence of frames "
    "IS the video.")
h2("9c. Finish & export")
bullet("Add captions/on-screen text and music. For any CLIP shots (Section 8d), drop the MP4 onto the timeline at "
       "that frame instead of drawing a still.")
bullet("Render/Download → MP4 at 1080p (or higher). That render is the finished episode.")
callout("PNG vs SVG: flat element PNGs reveal UNDER the hand (a wipe-on); SVG line versions (Section 7c) draw as true "
        "pen strokes. Mix as needed — SVG for hero draws, PNG for everything else.")

# ---------------- 10. PROCEDURE ----------------
h1("10. Step-by-Step Procedure When a Row Is Called")
for x in [
    "Read the row: concept, specific aspect, chosen animal match (higher %), and flags.",
    "Write the benefit-led suspense HOOK (3a).",
    "Write the full clean narration script using the retention arc (3b) to target length (3c).",
    "Compute duration (words ÷ 2.5) and segment into 3–5s FRAMES; assign timestamps (5a).",
    "Build the SHOT LIST (Deliverable A): per frame — narration slice, one-line visual, ELEMENT breakdown, draw "
    "order, file names, type.",
    "Build the PROMPT PACK (Deliverable B, Section 6): consistency header + per-element prompts grouped by frame, "
    "restating characters verbatim, each asking for a single isolated element on a transparent background.",
    "Generate the ELEMENTS (Section 7): run each prompt on Seedream 4.5 (free in the web app), reuse Host/prop "
    "Elements + style key-frame, then remove background → save transparent PNGs into frames/frame_0NN/.",
    "For any CLIP shots, generate the element still then animate it (Section 8d).",
    "Assemble in VideoScribe (Section 9): draw each element on, frame by frame; add VO/captions/music; render MP4.",
    "Save all deliverables per Section 5b naming, then run the checklists (Section 11).",
]:
    num(x)

# ---------------- 11. CHECKLISTS ----------------
h1("11. Quality Checklists")
h2("Script")
for c in ["Hook is benefit-led/suspenseful and doesn't just name the concept.",
          "Arc: stakes → animal → parallel → where-it-breaks-down → payoff → button.",
          "A question/open loop ~every 20–30s; reads naturally aloud; hits target runtime."]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2("Shot list & timing")
for c in ["Every frame 3–5s; each frame lists its elements + draw order; timestamps sum to total duration.",
          "Numbers match across shot list, prompt pack, and frame/element files."]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2("Prompt pack")
for c in ["Consistency header present (compact style tag + character/prop bible + palette + settings).",
          "Every per-element prompt restates characters VERBATIM, asks for a single isolated element on a "
          "transparent background, and is under 3000 characters.",
          "Grouped by frame with draw order; numbered to match; seed noted."]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2("Element PNGs (each)")
for c in ["Brand palette only; tomato coral reserved for the one focal element of the frame.",
          "Bright clip-art look: soft shading + gentle highlights, thin edge accents; no heavy black outlines, no "
          "grain, no gradients, no 3D.",
          "Single isolated element, TRANSPARENT background, clean cutout (no halo), no ground shadow.",
          "Recurring characters identical to their canonical Element; named frame_0NN_elXX_<name>.png."]:
    doc.add_paragraph("☐  " + c, style="List Bullet")
h2("Assembly")
for c in ["16:9 cream canvas; clear focal subject per frame; bottom 15% caption-safe.",
          "Elements drawn on in the planned order; timing matches the VO; frames flow in sequence."]:
    doc.add_paragraph("☐  " + c, style="List Bullet")

# ---------------- APPENDIX ----------------
h1("Appendix — Copy-Paste Blocks")
box("FULL MASTER STYLE SUFFIX (for single hero elements where length is not a constraint):", FULL_SUFFIX)
box("COMPACT STYLE TAG (for per-element prompts — keeps each under 3000 chars):", COMPACT_SUFFIX)
box("NEGATIVE / AVOID:",
    "scene background, environment, ground shadow, multiple subjects, photorealism, 3D render, CGI, realistic "
    "textures, film grain, gradients banding, glossy/metallic highlights, ambient occlusion, drop shadows, lens "
    "flare, glow/bloom, heavy black outlines, comic linework, sketchy wobble, crosshatching, dark/gritty mood, "
    "cluttered detail, complex detailed faces, tiny illegible text, watermarks, extra limbs/fingers, natural "
    "photographic animal coloring.")

note = doc.add_paragraph()
r = note.add_run("End of master guide. Input: one row of finance_animal_behavior_matches.xlsx. Output: a shot-list "
                 ".docx, a Higgsfield prompt pack, and a frames folder of transparent element PNGs — assembled in "
                 "VideoScribe with the animated hand-reveal, fixed series style, visuals every 3–5 seconds.")
r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor.from_string("888888")

out = "/home/user/dataexploration/PROJECT_MASTER_GUIDE.docx"
doc.save(out)
print("Saved:", out)
