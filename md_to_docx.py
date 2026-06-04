#!/usr/bin/env python3
"""Convert the Wonga Coup markdown dossier into a styled Word (.docx) document."""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "wonga_coup_dossier.md"
OUT = "wonga_coup_dossier.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREEN = RGBColor(0x2E, 0x7D, 0x32)   # FACT
AMBER = RGBColor(0xB7, 0x6E, 0x00)   # ALLEGED
RED = RGBColor(0xB0, 0x2A, 0x2A)     # DISPUTED
GREY = RGBColor(0x55, 0x55, 0x55)

TAG_COLORS = {"[FACT]": GREEN, "[ALLEGED]": AMBER, "[DISPUTED]": RED}

# Inline token regex: bold, italic, links, and the status tags.
INLINE = re.compile(
    r"(\*\*.+?\*\*|\[(?:FACT|ALLEGED|DISPUTED)\]|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)"
)


def add_runs(paragraph, text):
    """Render inline markdown (bold, italic, links, status tags) into runs."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part in TAG_COLORS:
            r = paragraph.add_run(part)
            r.bold = True
            r.font.color.rgb = TAG_COLORS[part]
        elif part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif re.fullmatch(r"\[[^\]]+\]\([^)]+\)", part):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            label, url = m.group(1), m.group(2)
            r = paragraph.add_run(label)
            r.font.color.rgb = NAVY
            r.underline = True
            # store the URL inline so it is not lost
            paragraph.add_run(f" ⟨{url}⟩").font.size = Pt(7)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def parse_table_row(line):
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# Heading colors
for lvl, color, size in [(1, NAVY, 20), (2, NAVY, 15), (3, NAVY, 12.5)]:
    st = doc.styles[f"Heading {lvl}"]
    st.font.color.rgb = color
    st.font.size = Pt(size)

with open(SRC, encoding="utf-8") as f:
    lines = f.read().split("\n")

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    stripped = line.strip()

    # --- Horizontal rule ---
    if stripped == "---":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("_" * 60)
        r.font.color.rgb = GREY
        i += 1
        continue

    # --- Headings ---
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
        i += 1
        continue
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        i += 1
        continue
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        i += 1
        continue

    # --- Tables ---
    if stripped.startswith("|") and i + 1 < n and set(lines[i + 1].strip()) <= set("|-: "):
        header = parse_table_row(line)
        body = []
        j = i + 2
        while j < n and lines[j].strip().startswith("|"):
            body.append(parse_table_row(lines[j]))
            j += 1
        table = doc.add_table(rows=1, cols=len(header))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, txt in enumerate(header):
            cell = table.rows[0].cells[c]
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], txt)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for row in body:
            cells = table.add_row().cells
            for c, txt in enumerate(row):
                if c < len(cells):
                    cells[c].paragraphs[0].text = ""
                    add_runs(cells[c].paragraphs[0], txt)
        doc.add_paragraph()
        i = j
        continue

    # --- Blockquote ---
    if stripped.startswith(">"):
        text = stripped.lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        if text:
            add_runs(p, text)
            for r in p.runs:
                r.italic = True
                if not r.font.color.rgb:
                    r.font.color.rgb = GREY
        i += 1
        continue

    # --- Bullet / numbered lists ---
    m_bullet = re.match(r"^(\s*)[-*]\s+(.*)", line)
    m_num = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
    if m_bullet:
        indent = len(m_bullet.group(1))
        style = "List Bullet" if indent < 2 else "List Bullet 2"
        p = doc.add_paragraph(style=style)
        add_runs(p, m_bullet.group(2))
        i += 1
        continue
    if m_num:
        p = doc.add_paragraph(style="List Number")
        add_runs(p, m_num.group(3))
        i += 1
        continue

    # --- Blank line ---
    if stripped == "":
        i += 1
        continue

    # --- Plain paragraph ---
    p = doc.add_paragraph()
    add_runs(p, stripped)
    i += 1

doc.save(OUT)
print(f"Wrote {OUT}")
