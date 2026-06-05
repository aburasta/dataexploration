#!/usr/bin/env python3
"""
Build the TIMED SHOT-LIST Word document for the Wonga Coup voiceover.

Segments the narration into ~3-5 second beats, assigns running timecodes
(at ~140 wpm), and prescribes one visual per beat from the researched asset
library, prioritizing real/archival over stock. Also writes a .md mirror.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

WPM = 140.0

# ----- colours -----
NAVY = RGBColor(0x1F, 0x3A, 0x5F)
TC = RGBColor(0x8A, 0x1C, 0x1C)      # timecode
TYPE_COLOR = {
    "REAL": RGBColor(0x2E, 0x7D, 0x32),         # green - real specific
    "CC": RGBColor(0x1F, 0x6F, 0xEB),           # blue - public domain / creative commons
    "ARCHIVE": RGBColor(0xB7, 0x6E, 0x00),      # amber - licensable news archive
    "STOCK": RGBColor(0x66, 0x66, 0x66),        # grey - generic stock
    "DOC": RGBColor(0x6A, 0x1B, 0x9A),          # purple - document/graphic
}

# ----- asset library (key -> (description, link, type, licence)) -----
A = {
    "tarmac_night": ("Night airport tarmac, parked jet under floodlights", "https://www.pexels.com/photo/black-private-jet-in-the-airport-14983522/", "STOCK", "Pexels, free, no attribution"),
    "harare_arrest": ("Detainees ('mercenaries') in custody, Harare court, March 2004 (agency news photo)", "https://www.aljazeera.com/news/2004/3/24/mercenaries-appear-in-harare-court", "ARCHIVE", "Underlying Reuters/AFP rights - license at source"),
    "b727": ("The actual Boeing 727, reg N4610 (MSN 727-18811) - airframe photos", "https://www.jetphotos.com/info/727-18811", "ARCHIVE", "Photographer-owned (JetPhotos) - request/license"),
    "detainees": ("The arrested men in khaki/shackles, Chikurubi special court 2004", "https://www.aljazeera.com/news/2004/3/24/mercenaries-appear-in-harare-court", "ARCHIVE", "Agency rights - license"),
    "soldiers": ("Soldier silhouette / armed men (generic)", "https://pixabay.com/photos/silhouette-soldier-military-917961/", "STOCK", "Pixabay, free, no attribution"),
    "merc32": ("32 'Buffalo' Battalion context (insignia/unit) + soldier B-roll", "https://en.wikipedia.org/wiki/32_Battalion_(South_Africa)", "CC", "Wikipedia/Commons - verify file licence"),
    "weapons": ("AK-47 / rifle close-up (generic)", "https://pixabay.com/photos/ak47-ak-47-gun-ak-47-rifle-3816651/", "STOCK", "Pixabay, free, no attribution"),
    "rpg": ("RPG / crates of weapons (generic)", "https://www.pexels.com/search/ak%2047/", "STOCK", "Pexels, free - pick & save specific file"),
    "eg_map": ("Equatorial Guinea location map (Bioko + Rio Muni)", "https://commons.wikimedia.org/wiki/File:Equatorial_Guinea_location_map.svg", "CC", "Wikimedia, free (SVG)"),
    "eg_globe": ("EG highlighted on the globe (Africa-centred) - establishing", "https://commons.wikimedia.org/wiki/File:Equatorial_Guinea_on_the_globe_(Africa_centered).svg", "CC", "Wikimedia, free (SVG)"),
    "obiang": ("President Teodoro Obiang Nguema - official portrait", "https://commons.wikimedia.org/wiki/Category:Teodoro_Obiang", "CC", "Wikimedia, free - confirm per-file tag"),
    "obiang_1979": ("Young Lt-Col Obiang just after the 1979 coup (moving footage)", "https://www.britishpathe.com/video/VLVA8KCYLX4E0E3ZQ0KYHDQJOQA7P-EQUATORIAL-GUINEA-PRESIDENT-TEODORO-OBIANG-NGUEMA-DESCRIBES/query/TOPPLING", "ARCHIVE", "British Pathe - license"),
    "obiang_recent": ("Obiang at a recent public/state appearance", "https://www.gettyimages.com/photos/teodoro-obiang-nguema-mbasogo", "ARCHIVE", "Getty - license (or Commons free)"),
    "cash": ("Money changing hands / stacks of cash (generic)", "https://www.pexels.com/photo/paper-money-changing-hands-6694571/", "STOCK", "Pexels, free"),
    "oil_rig": ("Offshore oil rig, Gulf of Guinea look (generic)", "https://www.pexels.com/photo/photo-of-oil-rig-platform-1716008/", "STOCK", "Pexels, free"),
    "oil_flare": ("Offshore rig / gas flare burning at night", "https://www.pexels.com/photo/illuminated-offshore-oil-rig-at-night-30475898/", "STOCK", "Pexels, free"),
    "zafiro": ("Zafiro field reference (the actual EG field, on-stream 1996)", "https://www.nsenergybusiness.com/projects/zafiro-oil-field-gulf-guinea/", "REAL", "Editorial reference - clear rights before reuse"),
    "malabo_streets": ("Real Malabo street scene", "https://commons.wikimedia.org/wiki/Category:Streets_in_Equatorial_Guinea", "CC", "Wikimedia - verify per-file licence"),
    "malabo_city": ("Modern Malabo skyline / cityscape", "https://www.gettyimages.com/photos/malabo", "ARCHIVE", "Getty - license (some Commons free)"),
    "malabo_slum": ("Malabo shanty/poor neighbourhood (wealth-gap contrast)", "https://commons.wikimedia.org/wiki/Category:Slums_in_Equatorial_Guinea", "CC", "Wikimedia - verify per-file licence"),
    "spanish_colonial": ("Colonial 'Spanish Guinea' period imagery", "https://en.wikipedia.org/wiki/Spanish_Guinea", "CC", "Most colonial-era images PD by age - verify"),
    "cocoa": ("Cocoa plantation, Fernando Po/Bioko (colonial era)", "https://stampworldhistory.nl/country-profiles-2/africa/fernando-poo/", "REAL", "Pre-1929 generally PD - verify"),
    "riggs": ("Riggs Bank (Washington DC) / US Senate investigation context", "https://en.wikipedia.org/wiki/Riggs_Bank", "CC", "Wikipedia/Commons + US-Gov Senate imagery (PD)"),
    "macias": ("Francisco Macias Nguema - rare portrait (c.1970)", "https://www.gettyimages.com/detail/news-photo/francisco-macias-nguema-the-first-president-of-equatorial-news-photo/96550539", "ARCHIVE", "Getty - license (one of very few images)"),
    "coup1979": ("1979 coup: putschists tearing down the Macias portrait, Mongomo", "https://www.gettyimages.com/photos/francisco-macias-nguema", "ARCHIVE", "Getty - license"),
    "mann": ("Simon Mann - portrait (post-release, c.2011)", "https://commons.wikimedia.org/wiki/Category:Simon_Mann", "CC", "Wikimedia, free - confirm per-file tag"),
    "mann_trial2008": ("Mann in the makeshift Malabo courtroom, June 2008", "https://www.gettyimages.com/detail/news-photo/british-mercenary-simon-mann-sits-in-a-hall-turned-news-photo/81642684", "ARCHIVE", "Getty - license (iconic 'prisoner' still)"),
    "mann_harare": ("Mann leaving court in Harare, March 2004", "https://www.gettyimages.com/photos/simon-mann", "ARCHIVE", "Getty - license"),
    "mann_footage": ("Reuters package - archival moving footage of Mann", "https://reuters.screenocean.com/record/2000388", "ARCHIVE", "Reuters/Screenocean - license"),
    "cricket": ("Cricket batsman (for Mann's England-captain family)", "https://pixabay.com/photos/cricket-sports-player-batsman-5365724/", "STOCK", "Pixabay, free"),
    "sandhurst": ("Royal Military Academy Sandhurst (real)", "https://commons.wikimedia.org/wiki/Category:Royal_Military_Academy_Sandhurst", "CC", "Wikimedia - verify per-file licence"),
    "moto": ("Severo Moto (VOA image, Toledo) - free", "https://commons.wikimedia.org/wiki/Category:Severo_Moto_Ns%C3%A1", "CC", "Wikimedia; VOA/US-gov likely PD - verify"),
    "moto_madrid": ("Moto presenting his 'government-in-exile', Madrid, Aug 2003", "https://www.gettyimages.com/detail/news-photo/severo-moto-nsa-president-of-the-progress-party-of-news-photo/2454459", "ARCHIVE", "Getty - license"),
    "thatcher": ("Sir Mark Thatcher outside the High Court, Cape Town 2004", "https://www.gettyimages.com/photos/sir-mark-thatcher-in-court-in-cape-town", "ARCHIVE", "Getty/AFP - license"),
    "margaret": ("Margaret Thatcher (context cutaway)", "https://commons.wikimedia.org/wiki/Category:Margaret_Thatcher", "CC", "Wikimedia - verify per-file licence"),
    "helicopter": ("Alouette-type helicopter (generic)", "https://www.pexels.com/search/helicopter/", "STOCK", "Pexels, free - pick specific file"),
    "calil": ("Ely Calil - SCARCE (camera-shy; ~one known photo). Plan fallback (silhouette/document).", "https://en.wikipedia.org/wiki/Ely_Calil", "ARCHIVE", "Likely uncleared - trace original rights-holder. GAP."),
    "dutoit": ("Nick du Toit - Malabo trial, handcuffed, Nov 2004", "https://www.gettyimages.co.uk/detail/news-photo/south-african-nick-du-toit-the-chief-accused-is-pictured-news-photo/51761171", "ARCHIVE", "Getty/AFP (D. Minkoh) - license"),
    "malabo_airport": ("Malabo airport / control tower (establishing)", "https://www.gettyimages.com/photos/malabo", "ARCHIVE", "Getty - license"),
    "zdi": ("Zimbabwe Defence Industries (state arms factory) context", "https://en.wikipedia.org/wiki/Zimbabwe_Defence_Industries", "CC", "Wikipedia/Commons - verify"),
    "blackbeach": ("Black Beach prison, Malabo (exterior)", "https://www.flickr.com/photos/wapster/2330635367", "CC", "Flickr - CHECK per-photo CC licence; attribute"),
    "wonga_note": ("The 'splodge of wonga' note - NO facsimile online; recreate as a prop/graphic from the published text", "https://en.wikipedia.org/wiki/2004_Equatorial_Guinea_coup_attempt", "DOC", "Recreate prop - GAP in archive"),
    "book": ("Adam Roberts, 'The Wonga Coup' - the definitive account", "https://profilebooks.com/work/the-wonga-coup/", "REAL", "Book cover - publisher permission for on-screen use"),
    "trawler": ("Fishing trawler (the 'fishermen who couldn't fish')", "https://www.pexels.com/search/fishing%20trawler/", "STOCK", "Pexels, free"),
    "london": ("Chatham House / London think-tank exterior", "https://en.wikipedia.org/wiki/Chatham_House", "CC", "Wikipedia/Commons - verify"),
    "bar": ("Dim bar interior (loose talk) - generic", "https://www.pexels.com/search/bar%20at%20night/", "STOCK", "Pexels, free"),
    "fi": ("Formula 1 / Grand Prix race car", "https://www.pexels.com/photo/a-formula-1-car-on-a-race-track-10807493/", "STOCK", "Pexels, free - editorial use; avoid implying sponsorship"),
    "phone": ("Unanswered telephone (generic)", "https://www.pexels.com/search/telephone/", "STOCK", "Pexels, free"),
    "document": ("Intelligence dossier / typed report (generic)", "https://www.pexels.com/search/documents/", "STOCK", "Pexels, free"),
    "mbeki": ("Thabo Mbeki, South African president (archival)", "https://www.gettyimages.com/photos/thabo-mbeki", "ARCHIVE", "Getty - license (or Commons free)"),
    "mugabe": ("Robert Mugabe / Zimbabwe government (archival)", "https://www.gettyimages.com/photos/robert-mugabe", "ARCHIVE", "Getty - license (or Commons free)"),
    "amnesty": ("Amnesty International report on the flawed Malabo trial", "https://www.amnesty.org/en/wp-content/uploads/2021/08/afr240052005en.pdf", "REAL", "(c) Amnesty - permission for on-screen use"),
    "courtroom": ("Courtroom / gavel (generic)", "https://www.pexels.com/photo/judge-in-courtroom-with-gavel-and-documents-34817075/", "STOCK", "Pexels, free"),
    "prison": ("Prison corridor / barbed wire (generic)", "https://pixabay.com/photos/prison-jail-barbed-wire-barbwire-482619/", "STOCK", "Pixabay, free"),
    "flag": ("Flag of Equatorial Guinea", "https://commons.wikimedia.org/wiki/File:Flag_of_Equatorial_Guinea.svg", "CC", "Wikimedia, free"),
    "spain": ("Spanish Congress of Deputies / Madrid (Moto's exile)", "https://commons.wikimedia.org/wiki/Category:Severo_Moto_Ns%C3%A1", "CC", "Wikimedia - verify"),
}

# ----- the beats: (text, asset_key) ; ("##", "SEGMENT NAME") for headers -----
BEATS = [
    ("##", "COLD OPEN"),
    ("On the seventh of March, 2004,", "tarmac_night"),
    ("a Boeing 727 sat on the tarmac at Harare airport, going nowhere.", "b727"),
    ("In and around it were sixty-odd men.", "detainees"),
    ("Most of them had, at one time or another, been soldiers -", "soldiers"),
    ("and several had been very good ones.", "merc32"),
    ("They had flown into Zimbabwe to collect a shopping list:", "weapons"),
    ("assault rifles, machine guns, a hundred rocket-propelled grenade launchers, three hundred hand grenades.", "rpg"),
    ("Enough to take a country.", "weapons"),
    ("The country was Equatorial Guinea, two thousand miles to the north-west.", "eg_map"),
    ("The plan was to fly there, overthrow the president,", "obiang"),
    ("and install a more agreeable one - in exchange for money, and oil.", "cash"),
    ("There was just one problem.", "tarmac_night"),
    ("Before the plane could be loaded, Zimbabwean soldiers surrounded it.", "harare_arrest"),
    ("The weapons never came.", "weapons"),
    ("And it would later emerge that the plot these men were about to be arrested for", "detainees"),
    ("had already been read - in writing, months earlier -", "document"),
    ("by at least three of the world's intelligence services.", "document"),
    ("This is the story of how a brewing heir, a Prime Minister's son,", "mann"),
    ("an exiled priest, and several dozen out-of-work soldiers", "moto"),
    ("tried to buy an African country -", "eg_globe"),
    ("and botched it so thoroughly that the whole affair went down in history", "detainees"),
    ("under a piece of British slang for cash.", "cash"),
    ("[TITLE CARD] THE WONGA COUP", "wonga_note"),

    ("##", "ACT ONE - THE PRIZE"),
    ("To understand why anyone would gamble their freedom on a place", "eg_map"),
    ("most people can't find on a map,", "eg_globe"),
    ("you have to start with what was underneath it.", "oil_rig"),
    ("Equatorial Guinea is small - about the size of Maryland,", "eg_map"),
    ("with, back then, well under a million people.", "malabo_streets"),
    ("It's the only country in Africa where the official language is Spanish,", "flag"),
    ("a leftover from the one sub-Saharan colony Spain ever really held onto.", "spanish_colonial"),
    ("For most of its history it was, in the cold language of trade, unimportant:", "spanish_colonial"),
    ("cocoa plantations on an island called Bioko, a quiet mainland, and not much else.", "cocoa"),
    ("Then, in 1995, an American oil company drilled off the coast", "oil_rig"),
    ("and hit the Zafiro field.", "zafiro"),
    ("Almost overnight, the numbers went mad.", "oil_flare"),
    ("Production climbed from nothing to a hundred thousand barrels a day,", "oil_rig"),
    ("then two hundred thousand, then more.", "oil_flare"),
    ("Income per head grew at something like forty percent a year.", "cash"),
    ("On paper, this tiny, half-forgotten country became one of the richest in Africa -", "malabo_city"),
    ("the so-called 'Kuwait of Africa.'", "oil_flare"),
    ("On paper.", "malabo_slum"),
    ("Because the money flowed almost entirely to one family.", "obiang"),
    ("A United States Senate investigation would later trace hundreds of millions of dollars", "riggs"),
    ("running through a single Washington bank -", "riggs"),
    ("much of it, the investigators found, controlled by the president, his son, and his nephew.", "obiang"),
    ("Meanwhile most Equatoguineans stayed exactly as poor as they had always been.", "malabo_slum"),
    ("The man at the centre of it was President Teodoro Obiang Nguema.", "obiang"),
    ("And he understood, better than most, how an African presidency could change hands -", "obiang_1979"),
    ("because he had taken his own at gunpoint.", "coup1979"),
    ("In 1979, Obiang - then a senior army officer -", "obiang_1979"),
    ("overthrew the country's first dictator: his own uncle, Francisco Macias Nguema.", "macias"),
    ("Macias had spent eleven years turning the country into a graveyard,", "macias"),
    ("killing or driving out something close to a third of the population -", "spanish_colonial"),
    ("a regime so brutal it earned the grim nickname 'the Auschwitz of Africa.'", "macias"),
    ("When it was over, Obiang had his uncle tried, convicted of genocide, and shot.", "coup1979"),
    ("So that was the prize, and that was the man guarding it:", "obiang"),
    ("a ruler sitting on an ocean of oil money,", "oil_flare"),
    ("atop a small, lightly defended state", "eg_map"),
    ("with a long habit of changing hands by force.", "coup1979"),
    ("To a certain kind of outsider, that didn't look like a fortress.", "malabo_city"),
    ("It looked like an opening.", "oil_flare"),

    ("##", "ACT TWO - THE MEN"),
    ("The man who saw the opening most clearly was named Simon Mann.", "mann"),
    ("If you were casting the part, you'd struggle to do better.", "mann"),
    ("Mann was born into the heart of the British establishment -", "sandhurst"),
    ("old brewing money,", "cash"),
    ("a father and a grandfather who had both captained the England cricket team.", "cricket"),
    ("Eton, then Sandhurst, then the SAS.", "sandhurst"),
    ("And when he left the army, he did the logical thing for a man with his skills and his appetites:", "soldiers"),
    ("he went into the business of private war,", "merc32"),
    ("co-founding mercenary firms that fought around the oil and diamonds of Angola and Sierra Leone.", "soldiers"),
    ("Mann would later insist he was only 'the manager, not the architect' of what came next.", "mann_trial2008"),
    ("Architect or not, he was the one who gathered the pieces.", "mann"),
    ("The first piece was a replacement president.", "moto"),
    ("His name was Severo Moto -", "moto"),
    ("a former Catholic priest turned opposition politician,", "moto"),
    ("who had broken with Obiang and fled to Spain,", "spain"),
    ("where he ran a government-in-exile and waited.", "moto_madrid"),
    ("The plan was simple: remove Obiang, fly Moto in, sit him in the chair.", "obiang"),
    ("In return - according to the prosecutors who later laid it all out -", "courtroom"),
    ("Moto's camp had offered the plotters something like 1.8 million dollars,", "cash"),
    ("and, more to the point, oil rights.", "oil_flare"),
    ("The second piece was money.", "cash"),
    ("And this is where the story acquires its celebrity.", "thatcher"),
    ("Among those accused of helping fund it was Sir Mark Thatcher -", "thatcher"),
    ("son of the former British Prime Minister -", "margaret"),
    ("who happened to be Mann's neighbour in Cape Town.", "thatcher"),
    ("Thatcher would eventually admit, in a South African court, to paying for a helicopter;", "thatcher"),
    ("he said he believed it was for an air-ambulance business,", "helicopter"),
    ("but conceded he'd had his doubts about what Mann was really up to.", "thatcher"),
    ("Another name that surfaced was a London oil tycoon, Ely Calil -", "calil"),
    ("whom Mann would later name in court as the real boss, 'the cardinal.'", "calil"),
    ("Calil denied everything. He was never convicted of anything.", "calil"),
    ("The third piece was muscle:", "merc32"),
    ("sixty or seventy hired men,", "detainees"),
    ("many of them veterans of South Africa's notorious 32 'Buffalo' Battalion -", "merc32"),
    ("hard, experienced, and, after apartheid, out of work.", "soldiers"),
    ("The fourth piece was logistics.", "eg_map"),
    ("An advance team, led by a South African arms dealer named Nick du Toit,", "dutoit"),
    ("was already on the ground in the capital, Malabo -", "malabo_city"),
    ("posing as businessmen, ready to seize the airport tower and talk the plane in.", "malabo_airport"),
    ("Mann bought a second-hand Boeing 727 from a company in Kansas to carry the men.", "b727"),
    ("And to arm them, he arranged to stop in Zimbabwe", "zdi"),
    ("and buy a small mountain of weapons from the state arms factory.", "zdi"),
    ("On paper - that word again - it was a tidy little plan.", "eg_map"),
    ("A plane, some guns, a puppet, and a payday.", "b727"),
    ("The trouble was the gap between paper and people.", "mann"),

    ("##", "ACT THREE - THE LEAK"),
    ("Here's the thing about a secret plot to overthrow a government:", "document"),
    ("it only works if it stays secret.", "document"),
    ("This one didn't. Not even slightly.", "bar"),
    ("The plotters, it turned out, could not stop talking.", "bar"),
    ("The journalist Adam Roberts, who wrote the definitive account of all this,", "book"),
    ("put it simply: nobody could keep their mouth shut.", "bar"),
    ("The scheme was reportedly aired at a London think tank.", "london"),
    ("Men talked in bars.", "bar"),
    ("And the cover stories were, to put it kindly, thin.", "document"),
    ("At one point the group hired two fishing trawlers -", "trawler"),
    ("crewed by men with essentially no experience of fishing, or of the sea.", "trawler"),
    ("To a former special-forces officer named Johann Smith,", "soldiers"),
    ("who was watching on behalf of President Obiang, this was not subtle.", "obiang"),
    ("Fishermen who can't fish are not fishermen.", "trawler"),
    ("Smith wrote it all up - in detail -", "document"),
    ("and sent warnings to British and American contacts in December and January,", "document"),
    ("months before the plane ever left the ground.", "b727"),
    ("From there the warning travelled the corridors of African power -", "eg_globe"),
    ("most likely up to South Africa's president,", "mbeki"),
    ("and from him across to Zimbabwe's.", "mugabe"),
    ("Which is why, when Mann's Boeing touched down in Harare to collect its guns,", "b727"),
    ("the authorities weren't surprised to see it.", "tarmac_night"),
    ("They were waiting.", "tarmac_night"),
    ("On the seventh of March, the trap closed.", "harare_arrest"),
    ("Mann and his men were arrested on the tarmac.", "harare_arrest"),
    ("The cover story - that they were just security guards bound for a mine in Congo -", "detainees"),
    ("did not survive contact with a single skeptical official.", "courtroom"),
    ("The next day, two thousand miles away in Malabo,", "eg_map"),
    ("du Toit and his advance party were rounded up too.", "dutoit"),
    ("The coup was over before it had fired a shot.", "tarmac_night"),
    ("And then it produced its most famous artefact.", "wonga_note"),
    ("From his cell in a Zimbabwean prison,", "prison"),
    ("Mann smuggled out a note to his wife and his lawyers.", "wonga_note"),
    ("In it he pleaded for help - and, above all, for cash:", "cash"),
    ("'It may be that getting us out comes down to a large splodge of wonga.'", "wonga_note"),
    ("He complained that two of his backers - whom he called 'Smelly' and 'Scratcher' -", "calil"),
    ("weren't returning calls.", "phone"),
    ("They had, he grumbled, asked the lawyers to ring back 'after the Grand Prix race was over.'", "fi"),
    ("The names, it was widely reported, were nicknames -", "thatcher"),
    ("'Scratcher' for Mark Thatcher, 'Smelly' for Ely Calil -", "thatcher"),
    ("though that decoding was never proven in a courtroom.", "courtroom"),
    ("It hardly mattered. The picture was indelible:", "mann_trial2008"),
    ("a coup falling apart, its organizer begging from a cell,", "prison"),
    ("while the money men were apparently too busy watching motor racing to pick up the phone.", "fi"),
    ("That one word - wonga - would name the whole disaster.", "cash"),

    ("##", "ACT FOUR - THE RECKONING"),
    ("What followed was years of grinding, lopsided justice.", "courtroom"),
    ("The hired men, mostly, served their time and went home.", "detainees"),
    ("Mann was convicted in Zimbabwe of trying to buy those weapons.", "mann_harare"),
    ("In South Africa, in January 2005, Mark Thatcher pleaded guilty", "thatcher"),
    ("to unwittingly helping finance the plot -", "thatcher"),
    ("and paid for it with a fine of around half a million dollars and a suspended sentence.", "cash"),
    ("Then he left the country.", "thatcher"),
    ("He has always maintained he never knowingly funded a coup.", "thatcher"),
    ("The harshest reckoning fell on the men who ended up inside Equatorial Guinea itself -", "blackbeach"),
    ("in Black Beach, the prison Obiang himself had once run.", "blackbeach"),
    ("Nick du Toit was sentenced to thirty-four years,", "dutoit"),
    ("and would later describe torture and long stretches of solitary confinement.", "prison"),
    ("A German member of the team died in custody within days;", "prison"),
    ("the government said malaria, while Amnesty International cited witnesses", "amnesty"),
    ("who said he had been beaten to death. The truth was never established.", "prison"),
    ("And in 2008, after a quiet, sudden extradition,", "mann_footage"),
    ("Simon Mann himself was flown to Black Beach.", "blackbeach"),
    ("In a Malabo courtroom he was sentenced to more than thirty-four years.", "mann_trial2008"),
    ("And there, the manager turned on his backers -", "mann_trial2008"),
    ("naming Calil and Thatcher from the witness stand.", "thatcher"),
    ("They rejected his account.", "calil"),
    ("It looked, for a while, as though Mann would die in that prison.", "blackbeach"),
    ("He didn't.", "mann"),
    ("In November 2009, in a final twist no screenwriter would dare invent,", "mann_footage"),
    ("President Obiang simply pardoned him - and du Toit too -", "obiang"),
    ("on humanitarian grounds.", "obiang"),
    ("After a little over a year, Simon Mann walked out of Black Beach", "mann"),
    ("and flew home to England:", "b727"),
    ("to the country house, and the comfortable life,", "mann"),
    ("that the whole scheme had been meant to gild.", "oil_flare"),

    ("##", "CLOSE"),
    ("In the end, the Wonga Coup changed almost nothing about Equatorial Guinea.", "malabo_city"),
    ("Obiang is still in power - the longest-serving president on earth.", "obiang_recent"),
    ("The oil still flows.", "oil_flare"),
    ("And the money still pools where it always did.", "cash"),
    ("What the plot really exposed wasn't a weakness in a small African state.", "eg_map"),
    ("It was a particular kind of arrogance -", "mann"),
    ("the old, half-colonial assumption that a country with enough oil and a weak enough army", "oil_flare"),
    ("was simply there for the taking,", "eg_globe"),
    ("by men with the right schools, the right contacts,", "sandhurst"),
    ("and a large enough splodge of wonga.", "cash"),
    ("They were wrong about the country.", "malabo_streets"),
    ("But they were right that the story was worth telling.", "oil_flare"),
]


def fmt(t):
    return f"{int(t // 60):02d}:{int(t % 60):02d}"


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for lvl, sz in [(1, 19), (2, 14)]:
        doc.styles[f"Heading {lvl}"].font.color.rgb = NAVY
        doc.styles[f"Heading {lvl}"].font.size = Pt(sz)

    doc.add_heading("THE WONGA COUP - Timed Voiceover & Shot List", level=1)
    meta = doc.add_paragraph()
    meta.add_run("Source: wonga_coup_script.md (Node 4)  |  ~13:19 runtime @140 wpm  |  "
                 "one visual per ~3-5s beat.\n").italic = True
    legend = doc.add_paragraph()
    legend.add_run("Asset legend:  ").bold = True
    for k, label in [("REAL", "real/specific"), ("CC", "public-domain/Creative Commons (free)"),
                     ("ARCHIVE", "licensable news archive"), ("STOCK", "generic stock"),
                     ("DOC", "document/graphic to create")]:
        r = legend.add_run(f"{k} ")
        r.bold = True
        r.font.color.rgb = TYPE_COLOR[k]
        legend.add_run(f"= {label}   ").font.size = Pt(9)
    disc = doc.add_paragraph()
    disc.add_run("Rights note: links are a vetted starting map for a picture/footage researcher, "
                 "NOT cleared rights. Confirm each licence at source; most 2004-08 event footage is "
                 "AFP/Reuters/AP/Getty and must be licensed. Priority used: real -> public-domain/CC "
                 "-> licensable archive -> stock.").italic = True
    doc.add_paragraph()

    # word-based running clock
    clock = 0.0
    md = ["# THE WONGA COUP - Timed Voiceover & Shot List\n",
          "_~13:19 @140 wpm. One visual per ~3-5s beat. Links = starting map, not cleared rights._\n"]

    for text, key in BEATS:
        if text == "##":
            doc.add_heading(key, level=2)
            md.append(f"\n## {key}\n")
            continue
        words = len(text.split())
        start, end = clock, clock + words / WPM * 60
        clock = end
        tcode = f"[{fmt(start)}-{fmt(end)}]"

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        rt = p.add_run(tcode + "  ")
        rt.bold = True
        rt.font.color.rgb = TC
        rt.font.size = Pt(9)
        p.add_run(text)

        desc, link, typ, lic = A[key]
        v = doc.add_paragraph()
        v.paragraph_format.left_indent = Inches(0.3)
        v.paragraph_format.space_after = Pt(6)
        vr = v.add_run("VISUAL: ")
        vr.bold = True
        vr.font.size = Pt(9)
        vr.font.color.rgb = NAVY
        v.add_run(desc + "  ").font.size = Pt(9)
        tr = v.add_run(f"[{typ}] ")
        tr.bold = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = TYPE_COLOR[typ]
        lr = v.add_run(f"{link}  ({lic})")
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        md.append(f"{tcode} {text}")
        md.append(f"   - VISUAL: {desc} [{typ}] {link} ({lic})\n")

    n = sum(1 for t, _ in BEATS if t != "##")
    doc.add_paragraph()
    doc.add_paragraph(f"Total beats: {n}  |  Computed runtime: {fmt(clock)}").italic = True
    doc.save("wonga_coup_shotlist.docx")
    open("wonga_coup_shotlist.md", "w").write("\n".join(md) + f"\n\n_Total beats: {n} | runtime {fmt(clock)}_\n")
    print(f"Beats: {n}  runtime {fmt(clock)}  -> wonga_coup_shotlist.docx / .md")


build()
