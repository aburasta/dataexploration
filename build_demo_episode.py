#!/usr/bin/env python3
"""Build the demo episode deliverables for:
Concept: Risk pooling  |  Animal: Vampire bat reciprocal blood sharing.
Outputs: EP07_shot-list.docx (Deliverable A) and EP07_prompt-pack.txt (Deliverable B).

New pipeline: each FRAME is decomposed into its individual ELEMENTS; every element is
generated as a SINGLE isolated clip-art asset on a TRANSPARENT background (Higgsfield ->
remove background), then the elements are drawn onto a 16:9 cream VideoScribe canvas by the
animated hand, in DRAW ORDER, frame by frame, to build the video."""

import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- fixed style + bibles (bright children's-book clip-art) ----------
STYLE_SUFFIX = (
    "STYLE: bright cheerful flat 2D vector children's-book clip-art; rounded cute friendly shapes, soft subtle "
    "shading + gentle highlights, thin minimal edge accents (no heavy black outlines); dot eyes, rosy cheeks; "
    "palette sunshine #FBC02D, sky #4FC3E8, leaf #7CC576, coral #F0654E, tangerine #F39237, blush #F6A9B6, stone "
    "#9DB0BC, brown #8D6E63, cream #FBF7EF, charcoal #3A3631; SINGLE isolated element, centered, TRANSPARENT "
    "background, no scene/no ground shadow; clean modern wholesome. No grain/gradient/3D/realistic/photoreal/heavy "
    "outline/clutter/scene-background."
)
ISO = "Single isolated element, centered, transparent background, no scene, no ground shadow."

HOST = "cute clip-art fox host (rounded chunky sky-blue body, charcoal dot eyes, rosy cheeks, tomato-coral knitted scarf)"
BAT = "cute clip-art vampire bat (rounded stone-gray body, sky-blue wing membranes, tiny charcoal dot eyes, rosy cheeks, small white fangs)"
FIG = "simple rounded clip-art person figure (sky-blue body, charcoal dot eyes, no other facial detail)"

# ---------- the episode: (act, narration, [ (element_name, element_prompt), ... in DRAW ORDER ], type) ----------
S = [
# ---- HOOK ----
("HOOK", "What if one disaster could never wipe you out —",
 [("host", f"{HOST}, leaning forward, intrigued, eyebrows raised, looking at camera")], "STILL"),
("HOOK", "no matter how unlucky you got?",
 [("host", f"{HOST}, head tilted with a questioning expression"),
  ("question_mark", "a single bold tomato-coral question-mark symbol")], "STILL"),
("HOOK", "Most people are always one bad month away from losing everything.",
 [("figure", f"{FIG}, standing small and alone, slightly worried"),
  ("crack", "a single thin jagged tomato-coral crack line (a fissure in the ground)")], "STILL"),
("HOOK", "But a tiny creature solved this millions of years before we invented money.",
 [("bat", f"{BAT}, perched small and friendly"),
  ("leaf", "a single large rounded leaf-green jungle leaf"),
  ("coin", "a single flat sunshine-yellow coin")], "STILL"),
("HOOK", "And once you see how it works,",
 [("host", f"{HOST}, one paw raised with a knowing look")], "STILL"),
("HOOK", "you'll never look at insurance the same way again.",
 [("host", f"{HOST}, gesturing to the side, friendly"),
  ("shield", "a single rounded tomato-coral shield icon (insurance)")], "STILL"),
# ---- STAKES ----
("STAKES", "Here's the uncomfortable truth.",
 [("host", f"{HOST}, serious, leaning in")], "STILL"),
("STAKES", "You can do everything right — save, work hard, plan —",
 [("figure", f"{FIG}, calm and tidy, standing"),
  ("coin_icon", "a small sky-blue coin icon"),
  ("tool_icon", "a small sky-blue wrench/tool icon"),
  ("checklist_icon", "a small sky-blue checklist clipboard icon")], "STILL"),
("STAKES", "and a single emergency can erase it all in a day.",
 [("coral_burst", "a sharp tomato-coral burst / star explosion shape"),
  ("cracked_icons", "three small cracked sky-blue icons (coin, tool, checklist) tipping over")], "STILL"),
("STAKES", "Relying on luck is a losing game.",
 [("figure", f"{FIG}, looking up, uneasy"),
  ("dice", "a single wobbling tomato-coral dice cube (a symbol of luck)")], "STILL"),
("STAKES", "So how do you make yourself disaster-proof?",
 [("host", f"{HOST}, palms turned up in a questioning shrug"),
  ("question_mark", "a single bold tomato-coral question-mark symbol")], "STILL"),
# ---- ANIMAL ----
("ANIMAL", "Meet the vampire bat.",
 [("bat", f"{BAT}, hanging upside down, friendly, looking at camera")], "STILL"),
("ANIMAL", "It survives on one thing: blood.",
 [("bat", f"{BAT}, hanging, friendly"),
  ("droplet", "a single glowing tomato-coral blood droplet")], "STILL"),
("ANIMAL", "And it lives on a knife's edge.",
 [("bat", f"{BAT}, balanced tensely, cautious posture"),
  ("branch", "a single thin warm-brown tree branch")], "STILL"),
("ANIMAL", "A vampire bat must feed almost every single night.",
 [("bat", f"{BAT}, hanging, alert"),
  ("moon", "a small tomato-coral crescent moon"),
  ("droplet", "a small tomato-coral blood droplet")], "STILL"),
("ANIMAL", "Go just two nights without a meal, and it starves.",
 [("bat_weak", "a weak pale-greyish clip-art vampire bat (washed-out stone-gray body, droopy posture, tiny charcoal dot eyes)"),
  ("two_moons", "two small charcoal crescent-moon icons"),
  ("hourglass", "a small charcoal hourglass icon")], "STILL"),
("ANIMAL", "But hunting is unreliable.",
 [("bat_flying", f"{BAT}, wings spread, flying and searching")], "STILL"),
("ANIMAL", "On any night, a healthy bat may find nothing at all.",
 [("bat_worried", f"{BAT}, landing with a small worried expression"),
  ("bare_leaf", "a single bare leaf-green leaf")], "STILL"),
# ---- PARALLEL ----
("PARALLEL", "So these bats made a deal.",
 [("bat_left", f"{BAT}, facing right"),
  ("bat_right", f"{BAT}, facing left")], "STILL"),
("PARALLEL", "A bat that feeds well throws up part of its meal",
 [("bat_fed", f"{BAT}, plump and well-fed with a soft tomato-coral glow at its belly, leaning to one side"),
  ("bat_empty", f"{BAT}, thin and empty, hopeful")], "STILL"),
("PARALLEL", "to share with a roost-mate that came home empty.",
 [("bat_giving", f"{BAT}, offering, leaning forward"),
  ("droplet", "a glowing tomato-coral blood droplet"),
  ("bat_receiving", f"{BAT}, mouth open, receiving")], "STILL"),
("PARALLEL", "It gives away blood it didn't need",
 [("bat_giving_close", f"{BAT}, close-up, calm and generous posture"),
  ("droplet", "a small tomato-coral blood droplet")], "STILL"),
("PARALLEL", "to keep a hungry neighbor alive.",
 [("bat_revived", f"{BAT}, brightening and revived, healthy stone-gray color, happy")], "STILL"),
("PARALLEL", "And here's the genius part: the bats remember.",
 [("host", f"{HOST}, tapping its head with a knowing smile"),
  ("lightbulb", "a small tomato-coral lightbulb icon")], "STILL"),
("PARALLEL", "The ones you feed are the ones who feed you",
 [("bat_cluster", "a small roost cluster of several cute clip-art vampire bats (stone-gray bodies, sky-blue wings, dot eyes)"),
  ("coral_lines", "a web of thin tomato-coral connecting arrows showing who fed whom")], "STILL"),
("PARALLEL", "when your own luck runs out.",
 [("dim_bat", f"{BAT}, dim and tired"),
  ("returning_droplet", "a glowing tomato-coral droplet traveling back along a lit coral arrow")], "STILL"),
("PARALLEL", "Each bat pays a small price on its good nights",
 [("bat", f"{BAT}, dropping something downward"),
  ("droplet", "a small tomato-coral blood droplet"),
  ("pool_jar", "a glowing shared jar of tomato-coral liquid (a pooled reserve)")], "STILL"),
("PARALLEL", "to buy protection on its bad ones.",
 [("bat", f"{BAT}, reaching and taking"),
  ("droplet_out", "a tomato-coral droplet being drawn upward out of a jar"),
  ("pool_jar", "a glowing shared jar of tomato-coral liquid (a pooled reserve)")], "STILL"),
("PARALLEL", "No single bat can survive its own bad luck alone.",
 [("lone_bat", f"{BAT}, tiny and overwhelmed, looking up"),
  ("dark_wave", "a large looming dark stone-gray wave shape (bad luck)")], "STILL"),
("PARALLEL", "But pooled across the colony, one bat's disaster",
 [("bat_group", "a cluster of several cute clip-art vampire bats together, looking up bravely"),
  ("coral_dome", "a soft protective tomato-coral dome / shield arc"),
  ("dark_wave", "a large looming dark stone-gray wave shape (bad luck)")], "STILL"),
("PARALLEL", "is covered by another bat's surplus.",
 [("coral_dome", "a soft protective tomato-coral dome / shield arc"),
  ("lifted_bat", f"{BAT}, lifted and glowing, revived"),
  ("surplus_glow", "a warm sunshine-yellow surplus glow")], "STILL"),
("PARALLEL", "That is insurance. Exactly.",
 [("bat_shield", "a bold graphic tomato-coral shield emblem formed out of many tiny clip-art bat silhouettes")], "STILL"),
("PARALLEL", "The premium is the blood you give when you're full.",
 [("bat", f"{BAT}, plump, dropping a droplet"),
  ("droplet_plus", "a tomato-coral blood droplet marked with a small charcoal '+'"),
  ("pool_jar", "a glowing shared jar of tomato-coral liquid (a pooled reserve)")], "STILL"),
("PARALLEL", "The payout is the blood you get when you're empty.",
 [("bat", f"{BAT}, thin, receiving"),
  ("droplet_arrow", "a tomato-coral blood droplet with a small charcoal payout arrow"),
  ("pool_jar", "a glowing shared jar of tomato-coral liquid (a pooled reserve)")], "STILL"),
("PARALLEL", "The colony is the insurance company.",
 [("bat_company", "a cute cluster of clip-art vampire bats grouped together like a friendly little company"),
  ("coral_sign", "a small rounded tomato-coral hanging sign")], "STILL"),
# ---- INSIGHT ----
("INSIGHT", "This is why insurance feels like a waste —",
 [("host", f"{HOST}, shrugging"),
  ("calendar", "a simple cream-and-charcoal wall calendar"),
  ("question_mark", "a small tomato-coral question-mark symbol")], "STILL"),
("INSIGHT", "until the day it doesn't.",
 [("host", f"{HOST}, patient expression, watching"),
  ("calendar_flip", "a wall calendar with rounded cream pages mid-flip (charcoal grid)")], "CLIP"),
("INSIGHT", "Every month you pay, and nothing happens.",
 [("figure", f"{FIG}, dropping a coin"),
  ("coin", "a small tomato-coral coin"),
  ("slot_box", "a simple sky-blue coin-slot box")], "CLIP"),
("INSIGHT", "You're the well-fed bat, giving away blood you didn't need.",
 [("figure_bat", f"a {FIG} with a faint friendly clip-art vampire-bat silhouette overlay"),
  ("droplet", "a small tomato-coral blood droplet")], "STILL"),
("INSIGHT", "That feeling of waste is the product working.",
 [("coin_to_shield", "a tomato-coral coin morphing into a rounded tomato-coral shield (a transformation)")], "STILL"),
("INSIGHT", "You're paying so that if disaster strikes, the colony catches you.",
 [("coral_bolt", "a sharp tomato-coral lightning bolt striking downward"),
  ("figure", f"{FIG}, safe and calm"),
  ("coral_dome", "a protective tomato-coral dome formed of small figures")], "STILL"),
# ---- BREAKDOWN ----
("BREAKDOWN", "But here's where the bats and your insurer part ways.",
 [("host", f"{HOST}, one eyebrow raised in a 'but...' gesture")], "STILL"),
("BREAKDOWN", "The bats run on memory and friendship.",
 [("bat_friends", "a warm group of cute clip-art vampire bats close together"),
  ("coral_lines", "soft tomato-coral friendship lines linking them")], "STILL"),
("BREAKDOWN", "A company runs on cold math.",
 [("abacus", "a precise stone-gray and charcoal abacus, hard-edged and tidy"),
  ("grid", "a cool charcoal geometric grid")], "STILL"),
("BREAKDOWN", "It has priced you precisely to stay profitable.",
 [("abacus", "a precise stone-gray and charcoal abacus, calculating"),
  ("price_tag", "a small precise tomato-coral price-figure tag")], "STILL"),
("BREAKDOWN", "And unlike a loyal roost-mate, it can deny your claim.",
 [("cold_hand", "a cold geometric stone-gray hand"),
  ("claim_folder", "a tomato-coral claim folder snapping shut")], "STILL"),
("BREAKDOWN", "So the bats teach you why pooling works.",
 [("host", f"{HOST}, nodding, holding out one paw"),
  ("bat_icon", "a small warm cute clip-art vampire-bat icon")], "STILL"),
("BREAKDOWN", "But you're trusting a spreadsheet — not a friend.",
 [("host", f"{HOST}, holding something up, comparing"),
  ("spreadsheet", "a flat stone-gray spreadsheet grid")], "STILL"),
# ---- LESSON ----
("LESSON", "So the lesson is simple.",
 [("host", f"{HOST}, confident, speaking to camera")], "STILL"),
("LESSON", "Don't try to survive every disaster alone.",
 [("straining_figure", f"a {FIG} straining under a huge heavy load by itself"),
  ("dark_load", "a huge heavy dark stone-gray load / boulder")], "STILL"),
("LESSON", "Pool the risks too big to carry by yourself.",
 [("figure_group", "several rounded clip-art person figures together, lifting easily"),
  ("load", "a large stone-gray load / boulder held up lightly")], "STILL"),
("LESSON", "Pay the small price on your good days,",
 [("figure", f"{FIG}, calmly placing a coin"),
  ("coin", "a small tomato-coral coin"),
  ("pool", "a glowing shared tomato-coral pool")], "STILL"),
("LESSON", "so a single bad one can never end the game.",
 [("coral_bolt", "a tomato-coral lightning bolt bouncing off harmlessly"),
  ("coral_dome", "a protective tomato-coral dome")], "STILL"),
# ---- BUTTON ----
("BUTTON", "A creature that drinks blood in the dark figured that out.",
 [("bat_wise", f"{BAT}, a wise look to camera"),
  ("droplet", "a single warm glowing tomato-coral droplet")], "STILL"),
("BUTTON", "The question is — will you?",
 [("host", f"{HOST}, pointing gently at the viewer with a friendly, challenging smile")], "STILL"),
("BUTTON", "See you in the next one.",
 [("host", f"{HOST}, giving a small warm wave")], "STILL"),
]

WPS = 2.5  # 150 wpm


def mmss(t):
    return f"{int(t)//60:02d}:{int(round(t))%60:02d}"


def slug(name):
    return re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")


# compute timings + element file names
rows = []
clock = 0.0
for i, (act, narr, elements, typ) in enumerate(S, start=1):
    wc = len(narr.split())
    dur = max(2.0, wc / WPS)
    start, end = clock, clock + dur
    els = []
    for j, (ename, edesc) in enumerate(elements, start=1):
        fname = f"frame_{i:03d}_el{j:02d}_{slug(ename)}.png"
        els.append(dict(j=j, name=ename, desc=edesc, file=fname))
    rows.append(dict(n=i, act=act, narr=narr, els=els, typ=typ,
                     wc=wc, start=start, end=end))
    clock = end
total = clock
total_words = sum(r["wc"] for r in rows)
total_elements = sum(len(r["els"]) for r in rows)
n_clips = sum(1 for r in rows if r["typ"] == "CLIP")

# one stable seed per recurring element kind (character/prop)
kinds = []
for r in rows:
    for e in r["els"]:
        if e["name"] not in kinds:
            kinds.append(e["name"])
seed_of = {k: 673346 + 137 * idx for idx, k in enumerate(kinds)}

OUT = "/home/user/dataexploration/EP07_risk-pooling_vampire-bat"
os.makedirs(os.path.join(OUT, "frames"), exist_ok=True)
os.makedirs(os.path.join(OUT, "videoscribe"), exist_ok=True)
os.makedirs(os.path.join(OUT, "clips"), exist_ok=True)
for r in rows:
    os.makedirs(os.path.join(OUT, "frames", f"frame_{r['n']:03d}"), exist_ok=True)


# ============================ Deliverable A: shot-list.docx ============================
def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexv)
    tcPr.append(shd)


def sct(cell, text, bold=False, size=8.5, color=None):
    cell.text = ""; r = cell.paragraphs[0].add_run(text)
    r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)


doc = Document()
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EPISODE 07 — SHOT LIST"); r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string("1F4E5F")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Why insurance feels like a scam — until the day it saves you"); r.italic = True; r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string("F0654E")

doc.add_heading("Episode meta", level=1)
for k, v in [
    ("Finance concept", "Risk pooling (individual catastrophic loss absorbed by the group's surplus)"),
    ("Animal match", "Vampire bat reciprocal blood sharing (quality 96%, web-verified)"),
    ("Target runtime", f"{mmss(total)} (pilot cut; extend to 10–15 min for full release)"),
    ("Word count", f"{total_words} words"),
    ("Narration pace", "150 wpm (2.5 words/sec)"),
    ("Frames / elements", f"{len(rows)} frames, {total_elements} element PNGs ({n_clips} clips)"),
    ("Pipeline", "Seedream 4.5 → remove background → transparent element PNGs → assembled in VideoScribe "
                 "(animated hand-reveal). Nano Banana Pro only for any text/number elements."),
]:
    pp = doc.add_paragraph(); pp.add_run(f"{k}: ").bold = True; pp.add_run(v)

doc.add_heading("Clean narration script (read top to bottom)", level=1)
last_act = None
for r in rows:
    if r["act"] != last_act:
        doc.add_paragraph()
        last_act = r["act"]
    para = doc.paragraphs[-1]
    if para.runs and not para.text.endswith(" ") and para.text:
        para.add_run(" ")
    para.add_run(r["narr"] + " ").font.size = Pt(11)

doc.add_heading("Frame list (each frame = its elements, drawn on in order)", level=1)
tbl = doc.add_table(rows=1, cols=6); tbl.style = "Table Grid"
hd = tbl.rows[0].cells
for i, h in enumerate(["Frame", "Time", "Narration", "Elements (draw order)", "Files (frames/…)", "Type"]):
    sct(hd[i], h, bold=True, size=8.5, color="FFFFFF"); shade(hd[i], "1F4E5F")
for r in rows:
    c = tbl.add_row().cells
    sct(c[0], f"{r['n']:03d}")
    sct(c[1], f"{mmss(r['start'])}-{mmss(r['end'])}")
    sct(c[2], r["narr"])
    sct(c[3], " → ".join(f"{e['j']}. {e['name']}" for e in r["els"]))
    sct(c[4], f"frame_{r['n']:03d}/ ({len(r['els'])} png" + (", +clip)" if r["typ"] == "CLIP" else ")"))
    sct(c[5], r["typ"])
widths = [0.45, 0.85, 2.1, 2.2, 1.1, 0.5]
for row in tbl.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = Inches(w)

doc.save(os.path.join(OUT, "EP07_shot-list.docx"))

# ============================ Deliverable B: prompt-pack.txt ============================
lines = []
lines.append("HIGGSFIELD PROMPT PACK  —  EPISODE 07: Risk pooling / Vampire bat reciprocal blood sharing")
lines.append("Model: Seedream 4.5  |  Output: SINGLE isolated element per prompt → remove background → transparent PNG")
lines.append("Each per-element prompt below is fully self-contained and under 3000 characters — paste it directly.")
lines.append("=" * 100)
lines.append("")
lines.append("=== CONSISTENCY HEADER  (keep these anchors identical for every element) ===")
lines.append("")
lines.append("RECURRING CHARACTERS (describe with these EXACT words whenever they appear):")
lines.append(f"  HOST  = {HOST}")
lines.append(f"  BAT   = {BAT}")
lines.append(f"  FIG   = {FIG}")
lines.append("")
lines.append("PALETTE (hex): sunshine #FBC02D, sky #4FC3E8, leaf #7CC576, coral #F0654E, tangerine #F39237, "
             "blush #F6A9B6, stone #9DB0BC, brown #8D6E63, cream #FBF7EF, charcoal #3A3631.")
lines.append("")
lines.append("HOW TO GENERATE EACH ELEMENT:")
lines.append("  - Model: Seedream 4.5 (free/unlimited in the web app)  |  quality: high for hero elements, basic otherwise")
lines.append("  - Generate ONE isolated element on a plain flat background, then REMOVE BACKGROUND to a transparent PNG")
lines.append("    (web-app remove-background, or MCP remove_background with media_id + media_type='image').")
lines.append("  - Reference image: attach your approved style key-frame so the look carries over.")
lines.append("  - Element/seed: reuse your saved 'host-fox' / 'bat' Element and the listed per-element seed so the same")
lines.append("    asset is identical everywhere it appears (don't regenerate a recurring asset — reuse the PNG).")
lines.append("  - Coral (#F0654E) is the focal/spotlight color: reserve it for the one focal element of each frame.")
lines.append("  - Save as frames/frame_0NN/frame_0NN_elXX_<name>.png.")
lines.append("")
lines.append("ASSEMBLY (VideoScribe): on a 16:9 CREAM (#FBF7EF) canvas, import each frame's element PNGs and let the")
lines.append("animated hand DRAW them on in the listed draw order across the frame's narration slice; then the next")
lines.append("frame; render MP4. For CLIP frames, generate the element still first, animate it, drop the clip in.")
lines.append("")
lines.append("STYLE BLOCK (appended to every element prompt below):")
lines.append("  " + STYLE_SUFFIX)
lines.append("")
lines.append("=" * 100)
lines.append("PER-FRAME / PER-ELEMENT PROMPTS")
lines.append("=" * 100)
for r in rows:
    lines.append("")
    order = " → ".join(f"{e['j']}.{e['name']}" for e in r["els"])
    lines.append(f"=== FRAME {r['n']:03d}  |  {mmss(r['start'])}-{mmss(r['end'])}  |  {r['typ']}  |  draw order: {order} ===")
    lines.append(f"    narration: \"{r['narr']}\"")
    if r["typ"] == "CLIP":
        lines.append("    [CLIP: generate these element stills first, then animate the moving one with image-to-video, same style.]")
    for e in r["els"]:
        lines.append("")
        lines.append(f"  -- el{e['j']:02d}  {e['name']}  (seed {seed_of[e['name']]})  ->  {e['file']}")
        lines.append("     " + e["desc"] + ". " + ISO + " " + STYLE_SUFFIX)
lines.append("")
lines.append("=" * 100)
lines.append(f"END OF PROMPT PACK  —  {len(rows)} frames, {total_elements} element prompts, runtime {mmss(total)}.")

with open(os.path.join(OUT, "EP07_prompt-pack.txt"), "w") as f:
    f.write("\n".join(lines))

prompt_lens = [len(e["desc"] + ". " + ISO + " " + STYLE_SUFFIX) for r in rows for e in r["els"]]
assert max(prompt_lens) < 3000, f"A prompt exceeds 3000 chars: {max(prompt_lens)}"

print("Wrote:")
print(" ", os.path.join(OUT, "EP07_shot-list.docx"))
print(" ", os.path.join(OUT, "EP07_prompt-pack.txt"))
print(f"Frames: {len(rows)} | elements: {total_elements} | words: {total_words} | runtime: {mmss(total)} | clips: {n_clips}")
print(f"Per-element prompt chars: min {min(prompt_lens)} | max {max(prompt_lens)} | style block {len(STYLE_SUFFIX)} (all < 3000)")
