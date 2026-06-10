#!/usr/bin/env python3
"""Generate the cartoon style guide (Word .docx) for the
finance x animal-behavior documentary series."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- helpers ----------
def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def h2(doc, text):
    return doc.add_heading(text, level=2)


def body(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def numbered(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def prompt_box(doc, label, text):
    """A shaded, monospace-ish copy-paste block."""
    if label:
        lp = doc.add_paragraph()
        r = lp.add_run(label)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string("1F4E5F")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F2F4F0")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph()


doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ============================ TITLE ============================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("VISUAL STYLE GUIDE")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = RGBColor.from_string("1F4E5F")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Finance × Animal Behavior — Animated Documentary Series")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string("EF5A2A")

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("The single source of truth for the look of every image and clip. "
                 "This is the art-direction module of the larger production guide.")
r.italic = True
r.font.size = Pt(10)

doc.add_paragraph()

# ============================ 0. HOW TO USE ============================
h1(doc, "0. How to Use This Guide")
body(doc, "This document defines the fixed visual language of the series so that hundreds of separately "
          "generated images stay perfectly on-brand and flow together as if hand-animated. Two audiences use it:")
bullet(doc, "A human (you) reviewing and approving images.")
bullet(doc, "An AI image/video model (Higgsfield) generating them. Every generation prompt MUST end with the "
            "Master Style Suffix in Section 13, and respect the Negative/Avoid list in Section 14.")
body(doc, "Golden rule: the STYLE never changes from video to video — only the CONTENT (which animal, which "
          "scene, which finance idea) changes. Consistency is the brand.")

# ============================ 1. NORTH STAR ============================
h1(doc, "1. The Look in One Sentence")
body(doc, "Bright, cheerful, wholesome FLAT 2D vector children's-book CLIP-ART — rounded cute friendly characters "
          "and objects, clean simple shapes, soft subtle shading with gentle highlights, thin minimal edge accents "
          "(no heavy black outlines), dot eyes and rosy cheeks. Every asset is a SINGLE isolated element on a "
          "transparent background — no scene, no setting, no ground shadow.")
body(doc, "Emotional target: warm, friendly, cheerful, and a little playful. Even when the topic is a market crash, "
          "the visuals stay bright and disarming — that contrast is part of the channel's charm.")

# ============================ 2. REFERENCE BREAKDOWN ============================
h1(doc, "2. Reference Image — What Defines the Style")
body(doc, "The approved reference is a set of bright, cute clip-art animals (toucan, monkey, zebra, giraffe, lion, "
          "elephant, rhino, crocodile) each isolated on a plain white background. It establishes these "
          "non-negotiable traits:")
bullet(doc, "CLIP-ART, ISOLATED: each asset is a SINGLE element, centered, on a TRANSPARENT background — no scene, no "
            "setting, no ground shadow. The reference shows one creature at a time, cleanly cut out.")
bullet(doc, "SHAPE-FIRST: characters and objects are built from clean geometric shapes (circles, ovals, rounded "
            "rectangles) with soft rounded corners. Chunky, stable, friendly.")
bullet(doc, "SOFT SHADING + HIGHLIGHTS: gentle subtle shading and soft highlights give a little depth — not flat-only, "
            "not 3D.")
bullet(doc, "THIN MINIMAL EDGE ACCENTS: forms are mostly separated by color; a thin darker edge accent appears only "
            "where shapes need separation — never a heavy black comic outline.")
bullet(doc, "BRIGHT MULTICOLOR PALETTE: sunshine yellow, sky blue, leaf green, tomato coral, tangerine, blush pink, "
            "stone gray, warm brown, cream, charcoal. Cheerful and saturated, but fixed.")
bullet(doc, "CLEAN, NO GRAIN: crisp clean vector edges; no film grain, no paper texture, no noise overlays.")
bullet(doc, "FRIENDLY FACES: small dot eyes and rosy cheek blush; simple happy expressions.")
bullet(doc, "GENEROUS NEGATIVE SPACE: one clear element, nothing cluttered around it.")

# ============================ 3. COLOR PALETTE ============================
h1(doc, "3. Color Palette")
body(doc, "Use these colors and (sparingly) tints/shades of them. A tight palette is what makes every frame feel like "
          "the same show. Hex values are the canonical reference.")

palette = [
    ("Sunshine Yellow", "#FBC02D", "Primary", "Beaks, manes, warm highlights."),
    ("Cream", "#FBF7EF", "Primary", "VideoScribe canvas / negative space."),
    ("Sky Blue", "#4FC3E8", "Secondary", "Water/sky elements, calm characters."),
    ("Leaf Green", "#7CC576", "Secondary", "Foliage, ground shapes."),
    ("Tomato Coral", "#F0654E", "Accent", "THE focal color — the key idea. Use sparingly."),
    ("Tangerine", "#F39237", "Accent", "Secondary warm accent (lion/monkey)."),
    ("Blush Pink", "#F6A9B6", "Accent", "Cheek blush, soft accents."),
    ("Stone Gray", "#9DB0BC", "Neutral", "Elephant/rhino bodies, neutral props."),
    ("Warm Brown", "#8D6E63", "Neutral", "Trunks, wood, branches."),
    ("Charcoal", "#3A3631", "Neutral", "Eyes / edge accents / fine detail — never pure black."),
]

tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
for i, t in enumerate(["Swatch", "Name", "Hex", "Role", "Usage"]):
    set_cell_text(hdr[i], t, bold=True, size=9, color="FFFFFF")
    shade_cell(hdr[i], "1F4E5F")
for name, hexv, role, usage in palette:
    cells = tbl.add_row().cells
    shade_cell(cells[0], hexv)
    set_cell_text(cells[0], "", size=9)
    set_cell_text(cells[1], name, bold=True, size=9)
    set_cell_text(cells[2], hexv, size=9)
    set_cell_text(cells[3], role, size=9)
    set_cell_text(cells[4], usage, size=9)
# widths
for row in tbl.rows:
    row.cells[0].width = Inches(0.6)
    row.cells[1].width = Inches(1.3)
    row.cells[2].width = Inches(0.9)
    row.cells[3].width = Inches(0.9)
    row.cells[4].width = Inches(3.2)

doc.add_paragraph()
body(doc, "Palette rules:")
bullet(doc, "Bright but balanced: spread the palette across elements so frames feel cheerful, not chaotic. Reserve "
            "TOMATO CORAL for the ONE focal element per frame — the key idea the viewer should look at.")
bullet(doc, "Cream is the canvas: it is the VideoScribe background / negative space, not an element fill.")
bullet(doc, "No pure black and no pure white. Use Charcoal for eyes and edge accents; cream for the canvas.")

# ============================ 4. SHAPE & FORM ============================
h1(doc, "4. Shape & Form Language")
bullet(doc, "Build everything from circles, ovals, and rounded rectangles. Corners are always softened.")
bullet(doc, "Chunky and stable: thick limbs, solid bodies, low centers of gravity. Nothing thin or spiky.")
bullet(doc, "Geometric simplification: a tree is a few rounded blobs; a building is rounded rectangles; a cloud is "
            "overlapping circles.")
bullet(doc, "Symmetry and balance over realism. Proportions are playful, not anatomically correct.")

# ============================ 5. LINE & TEXTURE ============================
h1(doc, "5. Line, Edge Accents & Texture")
bullet(doc, "Edge accents: allowed but minimal. Separate shapes with color first; where a shape genuinely needs "
            "separation, add a THIN edge accent in a darker tint of the same fill (or charcoal) — never a heavy black "
            "ink outline, never a comic line.")
bullet(doc, "No grain, no texture: clean flat color, NO film grain, NO paper texture, NO noise overlays anywhere.")
bullet(doc, "Edges: clean crisp vector edges. No sketchy, hand-drawn wobble; no crosshatching. Cutouts must be clean "
            "with no white halo (elements are isolated on transparent backgrounds).")

# ============================ 6. LIGHTING & SHADOW ============================
h1(doc, "6. Lighting & Shadow")
bullet(doc, "Soft subtle shading + gentle highlights: add a little depth to each element with soft tonal shading and "
            "gentle soft highlights (a slightly lighter top, a slightly darker base). Keep it subtle — friendly "
            "dimension, not realism.")
bullet(doc, "NO ground or cast shadows: every element is isolated on a transparent background, so there is no oval "
            "under the character and no shadow cast onto a surface.")
bullet(doc, "No 3D, no ambient occlusion, no glossy/metallic highlights, no glow/bloom, no lens flare.")

# ============================ 7. CHARACTER SYSTEM ============================
h1(doc, "7. Character Design System")
body(doc, "There are three character types. All obey the same shape, palette, and face rules.")

h2(doc, "7a. The Host (recurring narrator)")
body(doc, "A single friendly animal guide who appears in every episode for brand continuity. Recommended: a rounded "
          "SKY-BLUE fox with a coral scarf as a signature, small charcoal dot eyes and rosy cheeks, rendered as a "
          "single isolated clip-art element. Lock ONE design and reuse it everywhere as a saved Higgsfield ELEMENT.")
bullet(doc, "Keep a fixed turnaround: same proportions, same colors, same signature scarf in every appearance.")
bullet(doc, "Expression range via eyebrows + posture: curious, surprised, thinking, reassuring.")

h2(doc, "7b. The “Animal of the Week”")
body(doc, "Each episode features the animal from the chosen spreadsheet row (vampire bat, gazelle, honeybee, locust, "
          "etc.). Render it in-style: cute rounded clip-art form, RE-TINTED into the brand palette (not natural photo "
          "colors), soft subtle shading + highlights, dot eyes, rosy cheeks, friendly proportions — a single isolated "
          "element on a transparent background.")

h2(doc, "7c. Abstract / Human Figures (markets, crowds, “you”)")
body(doc, "For crowds, traders, or the viewer stand-in, use simple rounded clip-art figures in palette colors "
          "— no detailed faces. A crowd is a field of identical rounded shapes; the odd one out is coral.")

h2(doc, "7d. Faces & Expressions")
bullet(doc, "Eyes: small charcoal dots or short ovals. Optional tiny highlight catchlight.")
bullet(doc, "Mouths: minimal — a small line, oval, or curve. Often none.")
bullet(doc, "Rosy cheek blush (a soft blush-pink oval) for warmth; eyebrows do the emotional heavy lifting.")

h2(doc, "7e. Character Consistency (critical)")
bullet(doc, "Generate each character as an isolated transparent element and SAVE the Host (and any repeated prop) as a "
            "Higgsfield ELEMENT, then reuse it in every prompt so the character stays identical across frames. "
            "(See Section 12.)")
bullet(doc, "Keep one canonical reference image of each character and pass it as a reference in new generations.")
bullet(doc, "Always describe the character the same way, word-for-word, in prompts.")

# ============================ 8. ENVIRONMENTS ============================
h1(doc, "8. Environments & Backgrounds")
body(doc, "There are NO baked-in scene backgrounds anymore. Every asset is a single isolated element on a transparent "
          "background; the CREAM VideoScribe canvas IS the background. Sets become optional, not painted scenes.")
bullet(doc, "Backgrounds/sets become OPTIONAL simple isolated PROP elements — a lamp, a plant, a tree — each its own "
            "transparent PNG, drawn onto the cream canvas where it helps. Or omit them entirely for a clean look.")
bullet(doc, "Generous negative space: the cream canvas supports the subject; props never compete with the focal element.")
bullet(doc, "Recurring props (e.g. a lamp + plant the Host sits near) can be saved as ELEMENTS and reused for continuity, "
            "but they are placed onto the canvas, never composited into a single scene image.")
bullet(doc, "Depth via flat layered elements drawn on in order and slight soft shading, not perspective realism.")

# ============================ 9. FINANCE MOTIFS ============================
h1(doc, "9. Finance Motifs, Icons & Charts (in-style)")
body(doc, "Finance concepts must be drawn in the SAME clip-art style — no realistic screenshots or stock chart imagery. "
          "Each motif is its OWN isolated transparent element in the brand palette.")
bullet(doc, "Charts: bold rounded line/bar charts in palette colors; the key line is coral. Chunky, friendly, minimal axes.")
bullet(doc, "Coins / money: simple sunshine-yellow circles with a minimal symbol; stacks as rounded rectangles.")
bullet(doc, "Arrows / trends: thick rounded arrows; up = leaf green/positive, down = tangerine/coral for alarm.")
bullet(doc, "Abstractions: a “herd” as identical rounded shapes; “risk” as a wobbling shape; "
            "“growth” as a sprouting rounded plant. Keep one icon vocabulary across the series. Each is a single "
            "isolated element on a transparent background.")

# ============================ 10. TYPOGRAPHY ============================
h1(doc, "10. On-Screen Text & Typography")
bullet(doc, "Font: a friendly geometric rounded sans-serif (e.g., a Poppins / Quicksand / Nunito feel). One family only.")
bullet(doc, "Use text sparingly: key term, a number, a punchy phrase. Charcoal on light fields; cream on dark.")
bullet(doc, "Numbers and key terms can pop in coral for emphasis.")
bullet(doc, "Generating clean text inside images is unreliable — PREFER adding titles/captions in VideoScribe or your "
            "video editor. If a text element must be generated, use Nano Banana Pro (better text), keep it to 1–3 "
            "large words as its own isolated transparent element, and verify spelling.")

# ============================ 11. COMPOSITION ============================
h1(doc, "11. Composition & Framing")
bullet(doc, "Element PNGs have NO fixed aspect ratio — each is generated isolated and centered, then placed. Generate "
            "at the highest available resolution.")
bullet(doc, "The 16:9 LAYOUT is composed on the VideoScribe CREAM canvas (YouTube), where elements are positioned and "
            "scaled into the frame — not baked into the individual PNGs.")
bullet(doc, "One clear focal subject per frame, with breathing room on the canvas.")
bullet(doc, "Caption-safe zone: on the canvas, keep important content out of the bottom ~15% (room for "
            "subtitles/lower-thirds).")
bullet(doc, "Vary shot scale across a sequence (wide establishing → medium → close detail) to keep energy up.")

# ============================ 12. HIGGSFIELD PRODUCTION ============================
h1(doc, "12. Higgsfield Production Settings")
bullet(doc, "Image model: Seedream 4.5 is PRIMARY for the bulk of elements (unlimited/free in the web app on a paid "
            "plan). Use Nano Banana Pro only for elements containing text/numbers.")
bullet(doc, "Isolated element → transparent PNG: generate the single isolated element on a plain flat background, then "
            "run BACKGROUND REMOVAL (the web app's remove-background tool, or the MCP remove_background with the "
            "generation's media_id and media_type='image') to get the alpha cutout, and download as a transparent PNG.")
bullet(doc, "Consistency tools: save the Host and recurring props as ELEMENTS and reference them in every prompt; keep "
            "one approved STYLE KEY-FRAME and attach it as a reference so new elements inherit palette and finish. "
            "Reuse one seed per character/prop. This is the #1 lever for frames that flow together.")
bullet(doc, "Optional SVG line version: auto-trace an element in Illustrator/Inkscape so it draws as true pen strokes "
            "in VideoScribe; flat PNGs reveal under the hand instead.")
bullet(doc, "Short clips (optional): use image-to-video ONLY where motion is essential (a number ticking up, a chart "
            "drawing itself, a coin flip). Generate the element still first, then animate it. Keep clips 2–4s; "
            "default motion is the VideoScribe hand-reveal.")

# ============================ 13. MASTER PROMPT FORMULA ============================
h1(doc, "13. The Master Prompt Formula")
body(doc, "Every prompt describes ONE isolated element = [the ONE element described] + \"single isolated element, "
          "centered, transparent background, no scene, no ground shadow\" + [the Master Style Suffix]. "
          "Append the suffix verbatim to lock the look.")

prompt_box(doc, "PROMPT TEMPLATE:",
           "[The ONE isolated element — a character described identically each time, or a prop, or a motif] "
           "[with one clear pose/state], [where the coral accent goes if this is the focal element]. "
           "Single isolated element, centered, transparent background, no scene, no ground shadow. "
           "+ MASTER STYLE SUFFIX")

prompt_box(doc, "MASTER STYLE SUFFIX (paste at the end of EVERY element prompt):",
           "Bright, cheerful, wholesome flat 2D vector children's-book clip-art illustration. Rounded, cute, friendly "
           "characters and objects with clean simple shapes and chunky proportions. Soft subtle shading and gentle "
           "soft highlights for a little depth; thin minimal darker edge accents only where a shape needs separation "
           "(no heavy black outlines). Small dot eyes, rosy cheek blush, simple happy expressions. Bright multicolor "
           "brand palette: sunshine yellow (#FBC02D), sky blue (#4FC3E8), leaf green (#7CC576), tomato coral "
           "(#F0654E), tangerine (#F39237), blush pink (#F6A9B6), stone gray (#9DB0BC), warm brown (#8D6E63), cream "
           "(#FBF7EF), charcoal (#3A3631). SINGLE isolated element only, centered, on a plain TRANSPARENT background "
           "— no scene, no setting, no ground shadow. Clean, modern, wholesome. No grain, no gradients, no 3D, no "
           "realistic textures, no photorealism, no heavy black outlines, no clutter, no scene background.")

# ============================ 14. NEGATIVE / AVOID ============================
h1(doc, "14. Negative / Avoid List")
body(doc, "Steer away from (state these as things to avoid in prompts):")
avoid = ("scene background, environment, ground shadow, multiple subjects, photorealism, 3D render, CGI, realistic "
         "textures, film grain, gradients banding, glossy/metallic highlights, ambient occlusion, drop shadows, lens "
         "flare, glow/bloom, heavy black outlines, comic linework, sketchy wobble, crosshatching, dark/gritty mood, "
         "cluttered detail, complex detailed faces, tiny illegible text, watermarks, extra limbs/fingers, natural "
         "photographic animal coloring.")
prompt_box(doc, "NEGATIVE PROMPT / AVOID:", avoid)

# ============================ 15. WORKED EXAMPLES ============================
h1(doc, "15. Worked Example Prompts")
body(doc, "Each below is a SINGLE isolated element; append the Master Style Suffix when generating, then remove the "
          "background to a transparent PNG.")

prompt_box(doc, "A) The Host element:",
           "Single isolated clip-art Host fox — rounded friendly sky-blue body with a signature coral scarf, small "
           "charcoal dot eyes, rosy cheeks, one paw raised mid-explanation, warm curious expression, centered. The "
           "coral scarf is the focal accent. Single isolated element, centered, transparent background, no scene, no "
           "ground shadow. + MASTER STYLE SUFFIX")

prompt_box(doc, "B) Animal-of-the-week element (vampire bat):",
           "Single isolated clip-art vampire bat — rounded chunky body in stone gray with sky-blue wing membranes, "
           "tiny charcoal dot eyes, small white fangs, rosy cheeks, friendly, hanging upside down, centered. Single "
           "isolated element, centered, transparent background, no scene, no ground shadow. + MASTER STYLE SUFFIX")

prompt_box(doc, "C) Finance abstraction element (the odd-one-out):",
           "Single isolated clip-art rounded figure in tomato coral turning to look the opposite direction, cute and "
           "friendly, dot eyes, rosy cheeks, centered — the coral focal element of a herd. Single isolated element, "
           "centered, transparent background, no scene, no ground shadow. + MASTER STYLE SUFFIX")

prompt_box(doc, "D) Finance motif element (chart, optional clip):",
           "Single isolated clip-art chunky rounded line chart, a single bold tomato-coral line rising then dipping, "
           "minimal charcoal axes, no realistic UI, clean and friendly, centered — the coral line is the focal point. "
           "Single isolated element, centered, transparent background, no scene, no ground shadow. + MASTER STYLE SUFFIX")

# ============================ 16. SEQUENCE / FLOW ============================
h1(doc, "16. Making Frames Flow via the Animated Hand-Reveal")
body(doc, "Motion comes from the VideoScribe animated DRAWING HAND, not from composed camera moves. Each frame is a "
          "few isolated element PNGs drawn onto the cream canvas one-by-one, in a planned DRAW ORDER across that "
          "frame's narration slice — so the visual builds as the narrator speaks:")
bullet(doc, "Plan a DRAW ORDER per frame: list the element PNGs and the sequence the hand reveals them (e.g. Host "
            "first, then the prop, then the coral focal element last). The reveal IS the animation.")
bullet(doc, "Reuse the same Host / prop ELEMENTS across frames so recurring assets are identical and you don't "
            "regenerate them — drop the saved element back onto the canvas.")
bullet(doc, "PNG vs SVG draw behavior: flat element PNGs REVEAL under the moving hand (a wipe-on); an optional SVG line "
            "version of an element draws as TRUE pen strokes. Mix as needed — SVG for hero draws, PNG for the rest.")
bullet(doc, "Reserve a true short video CLIP only for motion the hand-reveal can't fake (counting numbers, a coin "
            "flip, fast movement); drop the MP4 onto the timeline at that frame.")

# ============================ 17. CHECKLIST ============================
h1(doc, "17. Per-Image Quality Checklist")
for item in [
    "Brand palette only, with tomato coral reserved for the one focal element?",
    "Clip-art look: soft subtle shading + gentle highlights, thin edge accents — no heavy black outlines, no grain?",
    "SINGLE isolated element on a clean TRANSPARENT background — no halo, no ground shadow, no scene?",
    "Dot eyes + rosy cheeks; cute rounded friendly form?",
    "Recurring characters identical to their canonical Element?",
    "16:9 layout + caption-safe bottom 15% handled on the VideoScribe canvas (not baked into the PNG)?",
    "Bright, friendly, uncluttered — on-brand mood?",
    "Any text limited and spelled correctly (or left for VideoScribe/the editor)?",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("☐  " + item)
    r.font.size = Pt(10.5)

doc.add_paragraph()
note = doc.add_paragraph()
r = note.add_run("This is the visual-style module. It will be merged into the overall production guide "
                 "(script structure, hook formula, image-timestamp syncing) per your earlier instructions.")
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor.from_string("888888")

out = "/home/user/dataexploration/style_guide.docx"
doc.save(out)
print("Saved:", out)
